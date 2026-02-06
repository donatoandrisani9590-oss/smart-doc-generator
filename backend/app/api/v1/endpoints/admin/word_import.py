"""
Word-Import Assistent API (v4.2 Feature: Kapitel 15.2.8)

Importiert bestehende Word-Vorlagen und konvertiert sie zu Klauseln.

Erweitert um Side-by-Side Mapping-Assistent (v4.2 Feature):
- Analysiert Dokument mit Position-Tracking
- Generiert HTML-Preview mit interaktiven Markern
- Unterstützt "Für alle ersetzen" Funktion
"""

import re
import io
import os
import tempfile
import logging
from datetime import datetime
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel, Field

from app.db import get_db
from app.models.documents import Clause
from app.services.document_analyzer import (
    DocumentAnalyzerWithPositions,
    SimilarityDetector,
    analyze_document,
    confirm_mappings,
    TEMP_DIR,
)
from app.services.placeholder_validator import (
    KNOWN_PLACEHOLDERS,
    get_all_placeholders,
    get_placeholders_by_category,
)

logger = logging.getLogger(__name__)

router = APIRouter()


# ══════════════════════════════════════════════════════════════════════════════
# SCHEMAS
# ══════════════════════════════════════════════════════════════════════════════

class ExtractedSection(BaseModel):
    """Ein erkannter Abschnitt aus dem Word-Dokument."""
    index: int
    title: str
    content: str
    content_html: str
    char_count: int
    has_existing_match: bool = False
    existing_clause_id: Optional[int] = None
    existing_clause_title: Optional[str] = None
    selected: bool = True  # Standardmäßig zum Import ausgewählt


class WordAnalysisResult(BaseModel):
    """Ergebnis der Word-Dokument-Analyse."""
    filename: str
    total_chars: int
    sections_count: int
    sections: List[ExtractedSection]
    warnings: List[str]
    headings_found: int
    tables_converted: int
    images_ignored: int


class ImportClauseRequest(BaseModel):
    """Request zum Importieren ausgewählter Abschnitte als Klauseln."""
    sections: List[ExtractedSection]
    country_code: str = "DE"
    category: str = "Importiert"


class ImportResult(BaseModel):
    """Ergebnis des Klausel-Imports."""
    imported_count: int
    skipped_count: int
    clause_ids: List[int]


# ══════════════════════════════════════════════════════════════════════════════
# WORD PARSING HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_content: bytes) -> dict:
    """
    Extrahiert Text und Struktur aus einem DOCX-Dokument.

    Returns:
        dict mit text, sections, warnings, stats
    """
    try:
        from docx import Document
        from docx.table import Table
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx ist nicht installiert. Bitte führen Sie 'pip install python-docx' aus."
        )

    doc = Document(io.BytesIO(file_content))

    full_text = []
    sections = []
    current_section = {"title": "", "content": [], "content_html": []}
    warnings = []
    stats = {
        "headings_found": 0,
        "tables_converted": 0,
        "images_ignored": 0,
    }

    # Paragraph-Patterns für Überschriften (§1, Art. 1, 1., etc.)
    heading_patterns = [
        r'^§\s*\d+',           # §1, § 1, §12
        r'^Art\.\s*\d+',       # Art. 1, Art. 12
        r'^\d+\.\s+[A-ZÄÖÜ]',  # 1. Titel, 2. Abschnitt
        r'^[IVXLCDM]+\.\s+',   # I. Römische Nummerierung
    ]

    def is_heading(text: str, style_name: str) -> bool:
        """Prüft ob ein Paragraph eine Überschrift ist."""
        # Word-Styles
        if style_name and 'heading' in style_name.lower():
            return True
        if style_name and 'überschrift' in style_name.lower():
            return True

        # Pattern-Matching
        text_stripped = text.strip()
        for pattern in heading_patterns:
            if re.match(pattern, text_stripped, re.IGNORECASE):
                return True

        return False

    def save_current_section():
        """Speichert den aktuellen Abschnitt."""
        if current_section["title"] or current_section["content"]:
            content_text = "\n".join(current_section["content"]).strip()
            content_html = "<br>".join(current_section["content_html"])

            if content_text:  # Nur nicht-leere Abschnitte
                sections.append({
                    "title": current_section["title"].strip() or f"Abschnitt {len(sections) + 1}",
                    "content": content_text,
                    "content_html": f"<p>{content_html}</p>" if content_html else "",
                })

    # Paragraphen durchgehen
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Ist das eine Überschrift?
        if is_heading(text, style_name):
            # Vorherigen Abschnitt speichern
            save_current_section()

            # Neuen Abschnitt starten
            current_section = {
                "title": text,
                "content": [],
                "content_html": [f"<strong>{text}</strong>"],
            }
            stats["headings_found"] += 1
        else:
            # Text zum aktuellen Abschnitt hinzufügen
            current_section["content"].append(text)

            # HTML-Formatierung
            html_text = text
            for run in para.runs:
                if run.bold:
                    html_text = html_text.replace(run.text, f"<strong>{run.text}</strong>")
                if run.italic:
                    html_text = html_text.replace(run.text, f"<em>{run.text}</em>")

            current_section["content_html"].append(html_text)

        full_text.append(text)

    # Letzten Abschnitt speichern
    save_current_section()

    # Tabellen verarbeiten
    for table in doc.tables:
        stats["tables_converted"] += 1
        warnings.append(f"Tabelle mit {len(table.rows)} Zeilen wurde zu Text konvertiert")

        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_text.append(row_text)

        if table_text:
            # Tabelle als eigenen Abschnitt oder zum letzten hinzufügen
            if sections:
                sections[-1]["content"] += "\n\n" + "\n".join(table_text)
            else:
                sections.append({
                    "title": "Tabelle",
                    "content": "\n".join(table_text),
                    "content_html": "<ul>" + "".join(f"<li>{t}</li>" for t in table_text) + "</ul>",
                })

    # Bilder zählen (werden ignoriert)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            stats["images_ignored"] += 1

    if stats["images_ignored"] > 0:
        warnings.append(f"{stats['images_ignored']} Bild(er) wurden ignoriert - bitte über Design-Manager hochladen")

    return {
        "text": "\n".join(full_text),
        "sections": sections,
        "warnings": warnings,
        "stats": stats,
    }


async def find_similar_clauses(
    db: AsyncSession,
    title: str,
    country_code: str,
) -> Optional[Clause]:
    """
    Sucht nach ähnlichen existierenden Klauseln.
    """
    # Exakte Titel-Übereinstimmung
    result = await db.execute(
        select(Clause).where(
            Clause.title.ilike(f"%{title}%"),
            Clause.country_code == country_code,
            Clause.is_active == True,
        ).limit(1)
    )
    return result.scalar_one_or_none()


# ══════════════════════════════════════════════════════════════════════════════
# API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/analyze", response_model=WordAnalysisResult)
async def analyze_word_document(
    file: UploadFile = File(...),
    country_code: str = "DE",
    db: AsyncSession = Depends(get_db),
):
    """
    Analysiert ein Word-Dokument und extrahiert die Struktur.

    Schritt 1 des Import-Wizards:
    - Lädt DOCX hoch
    - Extrahiert Text und erkennt Abschnitte
    - Prüft auf ähnliche existierende Klauseln
    """
    # Dateiformat prüfen
    if not file.filename.endswith(('.docx', '.DOCX')):
        raise HTTPException(
            status_code=400,
            detail="Nur .docx Dateien werden unterstützt. Ältere .doc Dateien bitte zuerst in Word als .docx speichern."
        )

    # Datei lesen
    content = await file.read()

    if len(content) > 10 * 1024 * 1024:  # 10 MB Limit
        raise HTTPException(
            status_code=400,
            detail="Datei zu groß. Maximum: 10 MB"
        )

    # Word-Dokument parsen
    try:
        result = extract_text_from_docx(content)
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Fehler beim Lesen der Datei: {str(e)}"
        )

    # Abschnitte in Schema konvertieren und Duplikate prüfen
    sections: List[ExtractedSection] = []

    for idx, section in enumerate(result["sections"]):
        # Nach ähnlicher Klausel suchen
        existing = await find_similar_clauses(db, section["title"], country_code)

        sections.append(ExtractedSection(
            index=idx,
            title=section["title"],
            content=section["content"],
            content_html=section["content_html"],
            char_count=len(section["content"]),
            has_existing_match=existing is not None,
            existing_clause_id=existing.id if existing else None,
            existing_clause_title=existing.title if existing else None,
            selected=existing is None,  # Nicht auswählen wenn Duplikat
        ))

    # Warnungen ergänzen
    warnings = result["warnings"]
    if not sections:
        warnings.append("Keine Abschnitte erkannt. Der gesamte Text wird als ein Abschnitt importiert.")

    duplicates = sum(1 for s in sections if s.has_existing_match)
    if duplicates > 0:
        warnings.append(f"{duplicates} Abschnitt(e) haben ähnliche existierende Klauseln")

    return WordAnalysisResult(
        filename=file.filename,
        total_chars=len(result["text"]),
        sections_count=len(sections),
        sections=sections,
        warnings=warnings,
        headings_found=result["stats"]["headings_found"],
        tables_converted=result["stats"]["tables_converted"],
        images_ignored=result["stats"]["images_ignored"],
    )


@router.post("/import", response_model=ImportResult)
async def import_clauses_from_word(
    request: ImportClauseRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Importiert ausgewählte Abschnitte als neue Klauseln.

    Schritt 2 des Import-Wizards:
    - Erstellt Klauseln aus den ausgewählten Abschnitten
    - Setzt Kategorie und Land
    """
    imported_ids = []
    skipped = 0

    for section in request.sections:
        if not section.selected:
            skipped += 1
            continue

        # Neue Klausel erstellen
        clause = Clause(
            title=section.title,
            content_html=section.content_html or f"<p>{section.content}</p>",
            country_code=request.country_code,
            category=request.category,
            version=1,
            is_active=True,
            approval_status="draft",  # Importierte Klauseln starten als Entwurf
        )

        db.add(clause)
        await db.flush()  # Um ID zu bekommen
        imported_ids.append(clause.id)

    await db.commit()

    return ImportResult(
        imported_count=len(imported_ids),
        skipped_count=skipped,
        clause_ids=imported_ids,
    )


@router.get("/templates")
async def get_import_tips():
    """
    Gibt Tipps für den Word-Import zurück.
    """
    return {
        "tips": [
            "Verwenden Sie Überschriften-Styles (Überschrift 1, Überschrift 2) für bessere Erkennung",
            "Paragraphen mit §, Art. oder Nummerierung werden automatisch erkannt",
            "Tabellen werden zu Listen konvertiert",
            "Bilder und Logos werden ignoriert - laden Sie diese über den Design-Manager hoch",
            "Platzhalter wie {{ name }} müssen nach dem Import manuell eingefügt werden",
        ],
        "supported_formats": [".docx"],
        "max_file_size_mb": 10,
        "heading_patterns": [
            "§1, §2, ... (Paragraphen)",
            "Art. 1, Art. 2, ... (Artikel)",
            "1., 2., ... (Nummeriert)",
            "I., II., ... (Römisch)",
            "Word Überschrift-Styles",
        ],
    }


# ══════════════════════════════════════════════════════════════════════════════
# SIDE-BY-SIDE MAPPING ASSISTENT - SCHEMAS
# ══════════════════════════════════════════════════════════════════════════════

class TextPositionSchema(BaseModel):
    """Position eines Wertes im Dokument."""
    paragraph_index: int
    run_index: int
    char_start: int
    char_end: int
    table_info: Optional[Dict[str, int]] = None


class DetectedValueSchema(BaseModel):
    """Ein erkannter potentieller Platzhalter-Wert."""
    id: str
    text: str
    value_type: str
    position: TextPositionSchema
    confidence: float
    suggested_placeholder: Optional[str] = None
    suggested_label: Optional[str] = None
    alternatives: List[str] = []
    is_confirmed: bool = False
    user_override: Optional[str] = None


class DocumentSectionSchema(BaseModel):
    """Ein logischer Abschnitt des Dokuments."""
    id: str
    section_type: str
    content: str
    html_content: str
    position_start: int
    position_end: int
    detected_values: List[DetectedValueSchema] = []


class AnalysisStatistics(BaseModel):
    """Statistiken der Dokumentanalyse."""
    total_sections: int
    total_detected_values: int
    values_by_type: Dict[str, int]
    average_confidence: float
    high_confidence_count: int
    suggested_placeholder_count: int


class MagicAnalysisResult(BaseModel):
    """Vollständiges Ergebnis der Magic-Analyse mit Position-Tracking."""
    document_id: str
    original_filename: str
    preview_html: str
    preview_pdf_path: Optional[str] = None
    sections: List[DocumentSectionSchema]
    detected_values: List[DetectedValueSchema]
    statistics: AnalysisStatistics
    created_at: str


class MappingConfirmation(BaseModel):
    """Bestätigung einer einzelnen Zuordnung."""
    value_id: str
    placeholder_name: str
    apply_to_similar: bool = False  # "Für alle ersetzen"


class ConfirmMappingsRequest(BaseModel):
    """Request zum Bestätigen mehrerer Zuordnungen."""
    document_id: str
    confirmations: List[MappingConfirmation]


class ConfirmMappingsResponse(BaseModel):
    """Antwort nach Bestätigung der Zuordnungen."""
    confirmed_count: int
    auto_applied_count: int  # Anzahl "Für alle ersetzen" Anwendungen
    remaining_unconfirmed: int
    placeholder_mapping: Dict[str, str]  # value_id -> placeholder_name


class SimilarValuesRequest(BaseModel):
    """Request zum Finden ähnlicher Werte."""
    document_id: str
    value_id: str
    threshold: float = 0.85


class SimilarValuesResponse(BaseModel):
    """Antwort mit ähnlichen Werten für "Für alle ersetzen"."""
    reference_value: DetectedValueSchema
    similar_values: List[DetectedValueSchema]
    total_similar: int


class GenerateClausesRequest(BaseModel):
    """Request zum Generieren von Klauseln aus dem Mapping."""
    document_id: str
    placeholder_mapping: Dict[str, str]  # value_id -> placeholder_name
    clause_title: str
    category: str = "Importiert"
    country_code: str = "DE"


class GenerateClausesResponse(BaseModel):
    """Antwort nach Generierung der Klauseln."""
    clause_id: int
    clause_title: str
    content_html: str
    placeholder_count: int


# In-Memory Storage für Analyse-Ergebnisse (in Produktion: Redis/DB)
_analysis_cache: Dict[str, Dict[str, Any]] = {}


# ══════════════════════════════════════════════════════════════════════════════
# SIDE-BY-SIDE MAPPING ASSISTENT - ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/magic/analyze", response_model=MagicAnalysisResult)
async def magic_analyze_document(
    file: UploadFile = File(...),
    generate_pdf: bool = False,
):
    """
    🪄 Magic-Analyse: Analysiert ein Word-Dokument mit Position-Tracking.

    Der erste Schritt des Side-by-Side Mapping-Assistenten:
    - Extrahiert Dokumentstruktur (Paragraphen, Tabellen, Listen)
    - Erkennt potentielle Platzhalter-Werte (Daten, Beträge, Namen, etc.)
    - Generiert HTML-Preview mit interaktiven Markern
    - Optional: PDF-Preview für exakte visuelle Darstellung

    Returns:
        MagicAnalysisResult mit allen erkannten Werten und Preview-HTML
    """
    # Dateiformat prüfen
    if not file.filename.endswith(('.docx', '.DOCX')):
        raise HTTPException(
            status_code=400,
            detail="Nur .docx Dateien werden unterstützt."
        )

    # Datei lesen und temporär speichern
    content = await file.read()

    if len(content) > 10 * 1024 * 1024:  # 10 MB Limit
        raise HTTPException(
            status_code=400,
            detail="Datei zu groß. Maximum: 10 MB"
        )

    # Temporäre Datei erstellen
    with tempfile.NamedTemporaryFile(
        suffix='.docx',
        delete=False,
        dir=str(TEMP_DIR)
    ) as tmp_file:
        tmp_file.write(content)
        tmp_path = tmp_file.name

    try:
        # Analyse durchführen
        result = analyze_document(
            docx_path=tmp_path,
            original_filename=file.filename,
            generate_pdf=generate_pdf
        )

        # In Cache speichern für spätere Verwendung
        _analysis_cache[result['document_id']] = result

        # Konvertieren zu Response-Schema
        return MagicAnalysisResult(
            document_id=result['document_id'],
            original_filename=result['original_filename'],
            preview_html=result['preview_html'],
            preview_pdf_path=result['preview_pdf_path'],
            sections=[
                DocumentSectionSchema(
                    id=s['id'],
                    section_type=s['section_type'],
                    content=s['content'],
                    html_content=s['html_content'],
                    position_start=s['position_start'],
                    position_end=s['position_end'],
                    detected_values=[
                        DetectedValueSchema(
                            id=v['id'],
                            text=v['text'],
                            value_type=v['value_type'],
                            position=TextPositionSchema(**v['position']),
                            confidence=v['confidence'],
                            suggested_placeholder=v['suggested_placeholder'],
                            suggested_label=v['suggested_label'],
                            alternatives=v['alternatives'],
                            is_confirmed=v['is_confirmed'],
                        )
                        for v in s['detected_values']
                    ]
                )
                for s in result['sections']
            ],
            detected_values=[
                DetectedValueSchema(
                    id=v['id'],
                    text=v['text'],
                    value_type=v['value_type'],
                    position=TextPositionSchema(**v['position']),
                    confidence=v['confidence'],
                    suggested_placeholder=v['suggested_placeholder'],
                    suggested_label=v['suggested_label'],
                    alternatives=v['alternatives'],
                    is_confirmed=v['is_confirmed'],
                )
                for v in result['detected_values']
            ],
            statistics=AnalysisStatistics(**result['statistics']),
            created_at=result['created_at'],
        )

    except Exception as e:
        logger.error(f"Magic analysis failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Analyse fehlgeschlagen: {str(e)}"
        )

    finally:
        # Temporäre DOCX-Datei löschen (PDF bleibt für Download)
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.get("/magic/preview/{document_id}")
async def get_preview_html(document_id: str):
    """
    Gibt das HTML-Preview für ein analysiertes Dokument zurück.

    Das HTML enthält interaktive Marker für erkannte Werte,
    die mit JavaScript für Hover/Click-Interaktionen verwendet werden.
    """
    if document_id not in _analysis_cache:
        raise HTTPException(
            status_code=404,
            detail="Dokument nicht gefunden. Bitte erneut analysieren."
        )

    return {
        "document_id": document_id,
        "preview_html": _analysis_cache[document_id]['preview_html'],
    }


@router.get("/magic/preview-pdf/{document_id}")
async def get_preview_pdf(document_id: str):
    """
    Gibt das PDF-Preview für ein analysiertes Dokument zurück.

    Das PDF zeigt das Originaldokument mit exakter Formatierung.
    """
    if document_id not in _analysis_cache:
        raise HTTPException(
            status_code=404,
            detail="Dokument nicht gefunden."
        )

    pdf_path = _analysis_cache[document_id].get('preview_pdf_path')
    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(
            status_code=404,
            detail="PDF-Preview nicht verfügbar. Analysieren Sie erneut mit generate_pdf=true."
        )

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=f"preview_{document_id}.pdf"
    )


@router.post("/magic/find-similar", response_model=SimilarValuesResponse)
async def find_similar_values(request: SimilarValuesRequest):
    """
    🔍 Findet ähnliche Werte für die "Für alle ersetzen" Funktion.

    Wenn ein Benutzer eine Zuordnung bestätigt (z.B. "01.01.2025" → eintrittsdatum),
    kann diese Funktion andere ähnliche Werte im Dokument finden,
    die automatisch dieselbe Zuordnung erhalten sollen.
    """
    if request.document_id not in _analysis_cache:
        raise HTTPException(
            status_code=404,
            detail="Dokument nicht gefunden."
        )

    cached = _analysis_cache[request.document_id]
    detected_values = cached['detected_values']

    # Referenzwert finden
    reference = None
    for v in detected_values:
        if v['id'] == request.value_id:
            reference = v
            break

    if not reference:
        raise HTTPException(
            status_code=404,
            detail=f"Wert {request.value_id} nicht gefunden."
        )

    # SimilarityDetector nutzen
    from app.services.document_analyzer import DetectedValue, TextPosition
    from dataclasses import asdict

    # Konvertiere zu DetectedValue Objekten
    detected_objs = []
    for v in detected_values:
        pos = TextPosition(**v['position'])
        detected_objs.append(DetectedValue(
            id=v['id'],
            text=v['text'],
            value_type=v['value_type'],
            position=pos,
            confidence=v['confidence'],
            suggested_placeholder=v.get('suggested_placeholder'),
            suggested_label=v.get('suggested_label'),
            alternatives=v.get('alternatives', []),
        ))

    ref_obj = next(d for d in detected_objs if d.id == request.value_id)

    detector = SimilarityDetector(detected_objs)
    similar = detector.find_similar(ref_obj, threshold=request.threshold)

    return SimilarValuesResponse(
        reference_value=DetectedValueSchema(
            id=reference['id'],
            text=reference['text'],
            value_type=reference['value_type'],
            position=TextPositionSchema(**reference['position']),
            confidence=reference['confidence'],
            suggested_placeholder=reference.get('suggested_placeholder'),
            suggested_label=reference.get('suggested_label'),
            alternatives=reference.get('alternatives', []),
            is_confirmed=reference.get('is_confirmed', False),
        ),
        similar_values=[
            DetectedValueSchema(
                id=s.id,
                text=s.text,
                value_type=s.value_type,
                position=TextPositionSchema(
                    paragraph_index=s.position.paragraph_index,
                    run_index=s.position.run_index,
                    char_start=s.position.char_start,
                    char_end=s.position.char_end,
                    table_info=s.position.table_info,
                ),
                confidence=s.confidence,
                suggested_placeholder=s.suggested_placeholder,
                suggested_label=s.suggested_label,
                alternatives=s.alternatives,
                is_confirmed=s.is_confirmed,
            )
            for s in similar
        ],
        total_similar=len(similar),
    )


@router.post("/magic/confirm", response_model=ConfirmMappingsResponse)
async def confirm_value_mappings(request: ConfirmMappingsRequest):
    """
    ✅ Bestätigt Zuordnungen zwischen erkannten Werten und Platzhaltern.

    Der zentrale Endpoint für den Side-by-Side Mapping-Assistenten:
    - Bestätigt einzelne Zuordnungen
    - Unterstützt "Für alle ersetzen" für ähnliche Werte
    - Aktualisiert den Cache mit bestätigten Mappings
    """
    if request.document_id not in _analysis_cache:
        raise HTTPException(
            status_code=404,
            detail="Dokument nicht gefunden."
        )

    cached = _analysis_cache[request.document_id]
    detected_values = cached['detected_values']

    confirmed_count = 0
    auto_applied_count = 0
    placeholder_mapping = {}

    # Mapping für schnellen Zugriff
    values_by_id = {v['id']: v for v in detected_values}

    for conf in request.confirmations:
        if conf.value_id not in values_by_id:
            continue

        # Direkte Bestätigung
        values_by_id[conf.value_id]['is_confirmed'] = True
        values_by_id[conf.value_id]['user_override'] = conf.placeholder_name
        placeholder_mapping[conf.value_id] = conf.placeholder_name
        confirmed_count += 1

        # "Für alle ersetzen" anwenden
        if conf.apply_to_similar:
            from app.services.document_analyzer import DetectedValue, TextPosition

            ref_value = values_by_id[conf.value_id]
            ref_obj = DetectedValue(
                id=ref_value['id'],
                text=ref_value['text'],
                value_type=ref_value['value_type'],
                position=TextPosition(**ref_value['position']),
                confidence=ref_value['confidence'],
            )

            # Alle Werte als DetectedValue Objekte
            all_objs = [
                DetectedValue(
                    id=v['id'],
                    text=v['text'],
                    value_type=v['value_type'],
                    position=TextPosition(**v['position']),
                    confidence=v['confidence'],
                )
                for v in detected_values
            ]

            detector = SimilarityDetector(all_objs)
            similar = detector.find_similar(ref_obj)

            for sim in similar:
                if sim.id not in placeholder_mapping:
                    values_by_id[sim.id]['is_confirmed'] = True
                    values_by_id[sim.id]['user_override'] = conf.placeholder_name
                    placeholder_mapping[sim.id] = conf.placeholder_name
                    auto_applied_count += 1

    # Cache aktualisieren
    _analysis_cache[request.document_id]['detected_values'] = detected_values

    # Zähle unbestätigte
    remaining = sum(1 for v in detected_values if not v.get('is_confirmed'))

    return ConfirmMappingsResponse(
        confirmed_count=confirmed_count,
        auto_applied_count=auto_applied_count,
        remaining_unconfirmed=remaining,
        placeholder_mapping=placeholder_mapping,
    )


@router.post("/magic/generate-clause", response_model=GenerateClausesResponse)
async def generate_clause_from_mapping(
    request: GenerateClausesRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    📝 Generiert eine neue Klausel aus dem bestätigten Mapping.

    Finaler Schritt des Magic-Imports:
    - Ersetzt erkannte Werte durch Platzhalter {{ placeholder_name }}
    - Erstellt eine neue Klausel in der Datenbank
    - Die Klausel kann sofort im Dokumentengenerator verwendet werden
    """
    if request.document_id not in _analysis_cache:
        raise HTTPException(
            status_code=404,
            detail="Dokument nicht gefunden."
        )

    cached = _analysis_cache[request.document_id]

    # HTML-Content aufbauen mit Platzhalter-Ersetzungen
    # Hier vereinfacht - in Produktion würde man den Original-HTML-Content
    # mit den Positionen durchgehen und die Ersetzungen vornehmen

    sections_html = []
    placeholder_count = 0

    for section in cached['sections']:
        content = section['content']

        # Werte durch Platzhalter ersetzen (in umgekehrter Reihenfolge für Positionserhalt)
        values_in_section = sorted(
            [v for v in section['detected_values'] if v['id'] in request.placeholder_mapping],
            key=lambda v: v['position']['char_start'],
            reverse=True
        )

        for value in values_in_section:
            placeholder = request.placeholder_mapping[value['id']]
            content = (
                content[:value['position']['char_start']] +
                f"{{{{ {placeholder} }}}}" +
                content[value['position']['char_end']:]
            )
            placeholder_count += 1

        sections_html.append(f"<p>{content}</p>")

    final_html = "\n".join(sections_html)

    # Klausel in DB erstellen
    clause = Clause(
        title=request.clause_title,
        content_html=final_html,
        country_code=request.country_code,
        category=request.category,
        version=1,
        is_active=True,
        approval_status="draft",
    )

    db.add(clause)
    await db.commit()
    await db.refresh(clause)

    # Cache aufräumen
    del _analysis_cache[request.document_id]

    return GenerateClausesResponse(
        clause_id=clause.id,
        clause_title=clause.title,
        content_html=final_html,
        placeholder_count=placeholder_count,
    )


@router.get("/magic/placeholders")
async def get_available_placeholders():
    """
    📋 Gibt alle verfügbaren System-Platzhalter zurück.

    Für die Dropdown-Auswahl im Mapping-Assistenten.
    Gruppiert nach Kategorie für bessere Übersicht.
    """
    return {
        "placeholders": get_all_placeholders(),
        "by_category": get_placeholders_by_category(),
    }


@router.delete("/magic/cache/{document_id}")
async def clear_analysis_cache(document_id: str):
    """
    🗑️ Löscht das gecachte Analyse-Ergebnis.

    Sollte aufgerufen werden wenn der Benutzer den Import abbricht
    oder die Seite verlässt.
    """
    if document_id in _analysis_cache:
        # Auch temporäre PDF löschen
        pdf_path = _analysis_cache[document_id].get('preview_pdf_path')
        if pdf_path and os.path.exists(pdf_path):
            try:
                os.unlink(pdf_path)
            except OSError:
                pass

        del _analysis_cache[document_id]
        return {"status": "deleted", "document_id": document_id}

    return {"status": "not_found", "document_id": document_id}
