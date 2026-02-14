"""
Document Upload API with AI Extraction

ContractHero-inspired feature: Upload existing documents and extract data.

Features:
- PDF/DOCX upload
- AI-powered metadata extraction (Mistral/Ollama - GDPR-compliant)
- Automatic field recognition
- Integration with document generator

Privacy: Uses Mistral (EU) or Ollama (local) - no US data transfer.
"""
import asyncio
import time

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import json
import logging
import os
import tempfile
import re
from datetime import datetime

logger = logging.getLogger(__name__)

from app.db import get_db
from app.api.deps import get_current_user
from app.models.documents import DocumentType
from app.services.llm_service import get_llm_service, LLMMessage, LLMConfig, log_llm_call

router = APIRouter()


# ═══════════════════════════════════════════════════════════════════════════
# SCHEMAS
# ═══════════════════════════════════════════════════════════════════════════

class ExtractedField(BaseModel):
    """A field extracted from the document."""
    field_name: str
    field_value: str
    confidence: float
    source_text: Optional[str] = None


class ExtractionResult(BaseModel):
    """Result of document analysis."""
    success: bool
    document_type_guess: Optional[str] = None
    document_type_confidence: float = 0.0
    extracted_fields: List[ExtractedField] = []
    raw_text_preview: Optional[str] = None
    page_count: int = 0
    file_type: str = ""
    provider: str = "unknown"
    error: Optional[str] = None


class ExtractionRequest(BaseModel):
    """Request for text extraction."""
    text: str
    document_type_id: Optional[int] = None
    language: str = "de"


# ═══════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from PDF file."""
    try:
        import PyPDF2

        text = ""
        page_count = 0

        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)

            for page in reader.pages:
                text += page.extract_text() + "\n"

        return text.strip(), page_count
    except ImportError:
        # Try pdfplumber as fallback
        try:
            import pdfplumber

            text = ""
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"

            return text.strip(), page_count
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="PDF-Extraktion nicht verfügbar. Bitte PyPDF2 oder pdfplumber installieren."
            )


def extract_text_from_docx(file_path: str) -> tuple[str, int]:
    """Extract text from DOCX file."""
    try:
        from docx import Document

        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        # Approximate page count (rough estimate based on paragraphs)
        page_count = max(1, len(doc.paragraphs) // 25)

        return text.strip(), page_count
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="DOCX-Extraktion nicht verfügbar. Bitte python-docx installieren."
        )


# ═══════════════════════════════════════════════════════════════════════════
# PATTERN-BASED EXTRACTION (Fast, no LLM)
# ═══════════════════════════════════════════════════════════════════════════

GERMAN_PATTERNS = {
    "vorname": [
        r"Vorname[:\s]+([A-ZÄÖÜ][a-zäöüß]+)",
        r"Herr\s+(\w+)\s+\w+",
        r"Frau\s+(\w+)\s+\w+",
    ],
    "nachname": [
        r"Nachname[:\s]+([A-ZÄÖÜ][a-zäöüß]+)",
        r"(?:Herr|Frau)\s+\w+\s+([A-ZÄÖÜ][a-zäöüß]+)",
        r"Familie\s+([A-ZÄÖÜ][a-zäöüß]+)",
    ],
    "gehalt": [
        r"(?:Brutto|Monats)?gehalt[:\s]+(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)\s*(?:€|EUR|Euro)?",
        r"(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)\s*(?:€|EUR|Euro)\s*(?:brutto|monatlich)",
        r"Vergütung[:\s]+(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)",
    ],
    "eintrittsdatum": [
        r"Eintrittsdatum[:\s]+(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
        r"Beginn[:\s]+(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
        r"ab\s+(?:dem\s+)?(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
    ],
    "position": [
        r"Position[:\s]+([^\n,]+)",
        r"(?:als|in der Position)\s+([^\n,]+)",
        r"Stellenbezeichnung[:\s]+([^\n,]+)",
    ],
    "wochenstunden": [
        r"(\d{1,2})\s*(?:Stunden|Std\.?)\s*(?:pro\s*Woche|wöchentlich)",
        r"Wochenstunden[:\s]+(\d{1,2})",
        r"Arbeitszeit[:\s]+(\d{1,2})",
    ],
    "urlaubstage": [
        r"(\d{1,2})\s*(?:Tage|Urlaubstage)\s*(?:Urlaub|jährlich)?",
        r"Urlaubsanspruch[:\s]+(\d{1,2})",
        r"Jahresurlaub[:\s]+(\d{1,2})",
    ],
}

DOCUMENT_TYPE_KEYWORDS = {
    "Arbeitsvertrag": ["arbeitsvertrag", "anstellung", "beschäftigung", "arbeitgeber", "arbeitnehmer"],
    "Kündigung": ["kündigung", "beendigung", "aufhebung", "kündigungsfrist"],
    "Zeugnis": ["zeugnis", "arbeitszeugnis", "bescheinigung", "leistung", "führung"],
    "Änderung": ["änderung", "nachtrag", "ergänzung", "anpassung"],
}


def extract_with_patterns(text: str) -> tuple[List[ExtractedField], str, float]:
    """Fast pattern-based extraction."""
    fields = []
    text_lower = text.lower()

    # Extract fields
    for field_name, patterns in GERMAN_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                fields.append(ExtractedField(
                    field_name=field_name,
                    field_value=value,
                    confidence=0.7,
                    source_text=match.group(0)[:100]
                ))
                break  # Take first match

    # Guess document type
    best_type = None
    best_score = 0

    for doc_type, keywords in DOCUMENT_TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > best_score:
            best_score = score
            best_type = doc_type

    confidence = min(best_score / 3, 1.0)  # Max confidence at 3+ keywords

    return fields, best_type or "Unbekannt", confidence


# ═══════════════════════════════════════════════════════════════════════════
# LLM-BASED EXTRACTION (Deep analysis)
# ═══════════════════════════════════════════════════════════════════════════

async def extract_with_llm(text: str, document_type_hint: Optional[str] = None, custom_instructions: str = "") -> tuple[List[ExtractedField], str, float, str]:
    """Deep extraction using LLM."""
    _llm_start = time.time()
    try:
        llm = await get_llm_service()

        # Truncate text if too long
        max_chars = 4000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n...[gekürzt]"

        # Build prompt
        prompt = f"""Analysiere diesen deutschen Vertragstext und extrahiere strukturierte Daten.

TEXT:
{text}

Extrahiere diese Felder falls vorhanden:
- vorname: Vorname der Person
- nachname: Nachname der Person
- position: Stellenbezeichnung/Position
- gehalt: Monatsgehalt (nur Zahl)
- eintrittsdatum: Startdatum (Format: TT.MM.JJJJ)
- wochenstunden: Arbeitszeit pro Woche (nur Zahl)
- urlaubstage: Urlaubstage pro Jahr (nur Zahl)
- probezeit: Probezeit in Monaten

Antworte im JSON-Format:
{{
    "document_type": "Arbeitsvertrag|Kündigung|Zeugnis|Änderung|Sonstiges",
    "confidence": 0.0-1.0,
    "fields": {{
        "feldname": "wert",
        ...
    }}
}}"""

        system_content = "Du bist ein Dokumentenanalyse-Experte für deutsche Personalverträge."
        if custom_instructions:
            system_content += custom_instructions

        messages = [
            LLMMessage(role="system", content=system_content),
            LLMMessage(role="user", content=prompt)
        ]

        _llm_start = time.time()
        result = await llm.chat_json(messages, LLMConfig(temperature=0.1, json_mode=True))
        asyncio.create_task(log_llm_call(
            feature="document_upload",
            latency_ms=int((time.time() - _llm_start) * 1000),
            user_id=None,
        ))

        if not isinstance(result, dict):
            return [], "Unbekannt", 0.0, "error"

        fields = []
        for field_name, value in result.get("fields", {}).items():
            if value:
                fields.append(ExtractedField(
                    field_name=field_name,
                    field_value=str(value),
                    confidence=0.9,
                ))

        doc_type = result.get("document_type", "Unbekannt")
        confidence = float(result.get("confidence", 0.5))
        provider = llm.provider

        return fields, doc_type, confidence, provider

    except Exception as e:
        asyncio.create_task(log_llm_call(
            feature="document_upload",
            latency_ms=int((time.time() - _llm_start) * 1000),
            error_message=str(e),
        ))
        logger.warning(f"LLM extraction failed: {e}")
        return [], "Unbekannt", 0.0, "error"


# ═══════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════

@router.post("/upload", response_model=ExtractionResult)
async def upload_and_extract(
    file: UploadFile = File(...),
    use_ai: bool = Form(True),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """
    Upload a document and extract metadata.

    Supports PDF and DOCX files.
    Uses pattern matching for fast extraction,
    optionally enhanced with AI (Mistral/Ollama).
    """
    # Validate file type by extension
    filename = file.filename or ""
    file_ext = filename.lower().split(".")[-1] if "." in filename else ""

    if file_ext not in ["pdf", "docx"]:
        raise HTTPException(
            status_code=400,
            detail="Nur PDF und DOCX Dateien werden unterstützt"
        )

    # Validate MIME type
    ALLOWED_MIME_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",  # Fallback for some clients
    }
    if file.content_type and file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Ungueltiger MIME-Type: {file.content_type}. Nur PDF und DOCX erlaubt."
        )

    # Validate magic bytes (file signature)
    first_bytes = await file.read(8)
    await file.seek(0)  # Reset file position

    PDF_MAGIC = b"%PDF"
    DOCX_MAGIC = b"PK\x03\x04"  # ZIP/DOCX signature

    if file_ext == "pdf" and not first_bytes.startswith(PDF_MAGIC):
        raise HTTPException(
            status_code=400,
            detail="Datei ist keine gueltige PDF-Datei (ungueltige Dateisignatur)"
        )
    elif file_ext == "docx" and not first_bytes.startswith(DOCX_MAGIC):
        raise HTTPException(
            status_code=400,
            detail="Datei ist keine gueltige DOCX-Datei (ungueltige Dateisignatur)"
        )

    # Save to temp file using chunked streaming (prevents memory issues with large files)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            tmp_path = tmp.name
            # Stream in 1MB chunks to avoid loading entire file in memory
            chunk_size = 1024 * 1024  # 1MB
            total_size = 0
            while chunk := await file.read(chunk_size):
                total_size += len(chunk)
                if total_size > MAX_FILE_SIZE:
                    os.unlink(tmp_path)
                    raise HTTPException(
                        status_code=413,
                        detail=f"Datei zu gross. Maximum: {MAX_FILE_SIZE // (1024 * 1024)} MB"
                    )
                tmp.write(chunk)
    except Exception as e:
        # Cleanup if temp file was created but writing failed
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(
            status_code=500,
            detail=f"Fehler beim Speichern der Datei: {str(e)}"
        )

    try:
        # Extract text
        if file_ext == "pdf":
            text, page_count = extract_text_from_pdf(tmp_path)
        else:
            text, page_count = extract_text_from_docx(tmp_path)

        if not text.strip():
            return ExtractionResult(
                success=False,
                error="Kein Text im Dokument gefunden",
                file_type=file_ext,
                page_count=page_count,
            )

        # Pattern-based extraction (always)
        pattern_fields, doc_type, confidence = extract_with_patterns(text)
        provider = "pattern"

        # Load custom AI instructions
        from app.services.ai_instructions import get_ai_instructions
        country_code = getattr(current_user, "country_code", "DE") or "DE"
        custom_instructions = await get_ai_instructions(db, country_code)

        # AI enhancement (optional)
        if use_ai:
            ai_fields, ai_type, ai_confidence, ai_provider = await extract_with_llm(text, doc_type, custom_instructions)

            if ai_confidence > confidence:
                doc_type = ai_type
                confidence = ai_confidence

            # Merge fields (AI has higher confidence)
            field_map = {f.field_name: f for f in pattern_fields}
            for af in ai_fields:
                if af.field_name not in field_map or af.confidence > field_map[af.field_name].confidence:
                    field_map[af.field_name] = af

            pattern_fields = list(field_map.values())

            if ai_provider != "error":
                provider = ai_provider

        return ExtractionResult(
            success=True,
            document_type_guess=doc_type,
            document_type_confidence=confidence,
            extracted_fields=pattern_fields,
            raw_text_preview=text[:500] if len(text) > 500 else text,
            page_count=page_count,
            file_type=file_ext,
            provider=provider,
        )

    finally:
        # Cleanup
        os.unlink(tmp_path)


@router.post("/extract-text", response_model=ExtractionResult)
async def extract_from_text(
    request: ExtractionRequest,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """
    Extract metadata from raw text (for copy-paste use case).
    """
    text = request.text.strip()

    if not text:
        return ExtractionResult(
            success=False,
            error="Kein Text angegeben",
        )

    # Pattern extraction
    pattern_fields, doc_type, confidence = extract_with_patterns(text)

    # Load custom AI instructions
    from app.services.ai_instructions import get_ai_instructions
    country_code = getattr(current_user, "country_code", "DE") or "DE"
    custom_instructions = await get_ai_instructions(db, country_code)

    # AI enhancement
    ai_fields, ai_type, ai_confidence, provider = await extract_with_llm(text, custom_instructions=custom_instructions)

    if ai_confidence > confidence:
        doc_type = ai_type
        confidence = ai_confidence

    # Merge fields
    field_map = {f.field_name: f for f in pattern_fields}
    for af in ai_fields:
        if af.field_name not in field_map or af.confidence > field_map[af.field_name].confidence:
            field_map[af.field_name] = af

    return ExtractionResult(
        success=True,
        document_type_guess=doc_type,
        document_type_confidence=confidence,
        extracted_fields=list(field_map.values()),
        raw_text_preview=text[:500] if len(text) > 500 else text,
        file_type="text",
        provider=provider,
    )
