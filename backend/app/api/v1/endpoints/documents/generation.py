from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import Annotated, Optional
from pydantic import BaseModel, Field
from app.api import deps
from app.db import get_db
from app.services.template_engine import TemplateEngine
from app.services.preview import evaluate_condition
from app.services.html_converter import HtmlToDocxConverter
from app.models import core as models
from app.models.documents import DocumentType, Clause, DocumentTypeClause
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
from bs4 import BeautifulSoup
from datetime import datetime
import os
import logging
import re
import json

router = APIRouter()
engine = TemplateEngine()
logger = logging.getLogger(__name__)

# Output directory for generated documents
OUTPUT_DIR = Path(__file__).parent.parent.parent.parent.parent / "storage" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ═══════════════════════════════════════════════════════════════════════════════
# MONTH NAMES - Loaded dynamically from country configuration
# v3.1: Uses storage/config/countries.json instead of hardcoded values
# ═══════════════════════════════════════════════════════════════════════════════

def get_month_names(country_code: str) -> dict[int, str]:
    """
    Get localized month names from country configuration.

    Args:
        country_code: ISO 3166-1 alpha-2 code (e.g., "DE", "IT")

    Returns:
        Dictionary mapping month number (1-12) to localized name
    """
    try:
        from app.services.country_config import get_country
        country = get_country(country_code)
        return {i + 1: name for i, name in enumerate(country.month_names)}
    except Exception as e:
        logger.warning(f"Could not load months for {country_code}: {e}")
        # Fallback to German
        return {
            1: "Januar", 2: "Februar", 3: "März", 4: "April",
            5: "Mai", 6: "Juni", 7: "Juli", 8: "August",
            9: "September", 10: "Oktober", 11: "November", 12: "Dezember"
        }

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT TEXTS - Loaded dynamically from config
# v3.1: Uses storage/config/document-texts.json instead of hardcoded values
# ═══════════════════════════════════════════════════════════════════════════════

def get_document_texts_for_country(country_code: str) -> dict[str, str]:
    """
    Get localized document texts from configuration.

    Args:
        country_code: ISO 3166-1 alpha-2 code (e.g., "DE", "IT")

    Returns:
        Dictionary with localized texts (title, between, and, etc.)
    """
    try:
        from app.services.document_texts import get_document_texts
        return get_document_texts(country_code)
    except Exception as e:
        logger.warning(f"Could not load document texts for {country_code}: {e}")
        # Fallback to German
        return {
            "title": "Arbeitsvertrag",
            "between": "Zwischen",
            "and": "und",
            "employer_label": '– nachstehend "Arbeitgeber" genannt –',
            "employee_label": '– nachstehend "Arbeitnehmer" genannt –',
            "employee_prefix": "Herr/Frau",
            "signature_location": "Sulzberg",
            "date_placeholder": "______________",
        }

# Backwards compatibility function
def get_document_text(country_code: str) -> dict[str, str]:
    """Legacy wrapper for get_document_texts_for_country."""
    return get_document_texts_for_country(country_code)


def format_date_localized(date_str: str, country_code: str = "DE") -> str:
    """
    Convert ISO date to localized format based on country code.

    Uses country configuration for month names.
    Supports formats:
    - ISO: 2026-01-15
    - German: 15.01.2026
    - Italian: 15/01/2026

    Returns:
    - German: "15. Januar 2026"
    - Italian: "15 gennaio 2026"
    """
    if not date_str:
        return date_str
    try:
        if '-' in date_str:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
        elif '.' in date_str:
            # Check if already in German format (DD.MM.YYYY)
            parts = date_str.split('.')
            if len(parts) == 3:
                dt = datetime.strptime(date_str, "%d.%m.%Y")
            else:
                return date_str
        elif '/' in date_str:
            # Italian format (DD/MM/YYYY)
            dt = datetime.strptime(date_str, "%d/%m/%Y")
        else:
            return date_str

        # Get month names from country configuration
        months = get_month_names(country_code)
        month_name = months.get(dt.month, str(dt.month))

        if country_code == "IT":
            # Italian format: "15 gennaio 2026" (no dot after day)
            return f"{dt.day} {month_name} {dt.year}"
        else:
            # German format: "15. Januar 2026" (dot after day)
            return f"{dt.day:02d}. {month_name} {dt.year}"
    except (ValueError, KeyError) as e:
        logger.warning(f"Date formatting error: {e}")
        return date_str


def format_date_german(date_str: str) -> str:
    """Convert ISO date (2026-01-01) to German format (01. Januar 2026)."""
    return format_date_localized(date_str, "DE")


def render_placeholders_extended(content: str, form_data: dict, country_code: str = "DE") -> str:
    """
    Replace {{ placeholder }} and [placeholder] with formatted values.
    Extended version with localized date formatting.
    """
    def replace_match(match: re.Match) -> str:
        placeholder = match.group(1).strip()

        # Special handling for [datum] - use eintrittsdatum or current date
        if placeholder.lower() == "datum":
            value = form_data.get("eintrittsdatum") or form_data.get("datum")
            if value:
                return format_date_localized(str(value), country_code)
            # Fallback to current date
            return format_date_localized(datetime.now().strftime("%Y-%m-%d"), country_code)

        value = form_data.get(placeholder)

        if value is None:
            return f"[{placeholder}]"

        # Date fields - format with localization
        if "datum" in placeholder.lower() or "date" in placeholder.lower():
            return format_date_localized(str(value), country_code)

        # Currency fields
        if "gehalt" in placeholder.lower() or "betrag" in placeholder.lower():
            try:
                # Format as European currency
                num = float(value)
                formatted = f"{num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                return f"{formatted} €"
            except (ValueError, TypeError):
                return str(value)

        return str(value)

    # Replace both {{ placeholder }} and [placeholder] patterns
    content = re.sub(r"\{\{\s*(\w+)\s*\}\}", replace_match, content)
    content = re.sub(r"\[(\w+)\]", replace_match, content)
    return content


class GenerateDocumentRequest(BaseModel):
    """Request body for document generation."""
    # Required field - signatory must always be provided (min 2 characters)
    signatory_name: str = Field(
        ...,
        min_length=2,
        description="Pflichtfeld: Name des Unterzeichners (z.B. 'Matthias Schweizer')"
    )

    # Employee data
    vorname: Optional[str] = None
    nachname: Optional[str] = None
    geburtsdatum: Optional[str] = None
    eintrittsdatum: Optional[str] = None
    position: Optional[str] = None
    gehalt: Optional[str] = None
    wochenstunden: Optional[str] = None
    firmenwagen: Optional[bool] = None
    homeoffice: Optional[bool] = None
    probezeit: Optional[str] = None

    # Output options
    output_format: str = "pdf"
    attachment_ids: Optional[list[int]] = None
    clause_ids: Optional[list[int]] = None

    # Async processing option (recommended for PDF to avoid timeouts)
    async_pdf: bool = Field(
        default=False,
        description="If true, PDF conversion runs async and returns a job_id for polling"
    )

    class Config:
        extra = "allow"  # Erlaube zusätzliche Felder


class AsyncJobResponse(BaseModel):
    """Response for async PDF conversion jobs."""
    job_id: str
    status: str = "pending"
    message: str = "PDF conversion started. Poll /generate/jobs/{job_id} for status."


class JobStatusResponse(BaseModel):
    """Response for job status queries."""
    job_id: str
    status: str  # pending, processing, completed, failed
    output_path: Optional[str] = None
    error: Optional[str] = None
    download_url: Optional[str] = None


def html_to_docx_paragraph(doc: Document, html_content: str, form_data: dict, country_code: str = "DE"):
    """
    Convert HTML content to DOCX paragraphs with placeholder replacement.

    Handles:
    - <p> tags as paragraphs
    - <strong>/<b> as bold
    - <em>/<i> as italic
    - <br> as line breaks
    - {{placeholder}} replacement
    """
    # First, render placeholders with extended formatting
    rendered_html = render_placeholders_extended(html_content, form_data, country_code)

    # Parse HTML
    soup = BeautifulSoup(rendered_html, 'html.parser')

    # If no HTML tags, just add as plain text
    if not soup.find(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'div']):
        # Plain text - split by line breaks
        text = soup.get_text()
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        return

    # Process each element
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'div'], recursive=False):
        if element.name in ['h1', 'h2', 'h3']:
            # Headings
            para = doc.add_paragraph()
            run = para.add_run(element.get_text())
            run.bold = True
            if element.name == 'h1':
                run.font.size = Pt(16)
            elif element.name == 'h2':
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
        elif element.name in ['p', 'div']:
            # Regular paragraph
            para = doc.add_paragraph()

            # Process inline elements and text
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    if child.name in ['strong', 'b']:
                        run = para.add_run(child.get_text())
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run = para.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'br':
                        para.add_run('\n')
                    else:
                        text = child.get_text()
                        if text:
                            para.add_run(text)
                else:
                    # Text node
                    text = str(child)
                    if text and text.strip():
                        para.add_run(text)
        elif element.name in ['ul', 'ol']:
            # Lists
            for li in element.find_all('li', recursive=False):
                # Use bullet style for ul, number for ol
                try:
                    style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    para = doc.add_paragraph(style=style)
                except:
                    para = doc.add_paragraph()
                    prefix = "• " if element.name == 'ul' else ""
                    para.add_run(prefix)
                para.add_run(li.get_text())


def create_document_from_clauses(
    clauses: list[dict],
    form_data: dict,
    country_code: str,
    design_settings: dict,
    output_path: Path
) -> Path:
    """
    Create a DOCX document from assembled clauses.
    Professional format following DIN 5008 (DE) or UNI (IT) standards.
    Fully localized based on country_code.
    """
    doc = Document()

    # Get localized text from configuration
    loc_text = get_document_texts_for_country(country_code)

    # Set up page format - A4 with compliant margins (configurable)
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.left_margin = Cm(design_settings.get('margin_left_cm', 2.5))
    section.right_margin = Cm(design_settings.get('margin_right_cm', 2.0))
    section.top_margin = Cm(design_settings.get('margin_top_cm', 2.5))
    section.bottom_margin = Cm(design_settings.get('margin_bottom_cm', 2.0))
    section.header_distance = Cm(design_settings.get('header_distance_cm', 1.25))
    section.footer_distance = Cm(design_settings.get('footer_distance_cm', 1.0))

    # Set up document styles - configurable font size (default 11pt)
    font_size = design_settings.get('font_size_pt', 11)
    line_spacing = design_settings.get('line_spacing', 1.15)

    style = doc.styles['Normal']
    style.font.name = design_settings.get('font_family', 'Arial')
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing

    # Add logo to header (right-aligned) - professional size with spacing
    logo_path = design_settings.get('logo_path')
    logo_added = False

    # Try multiple possible logo paths
    possible_logo_paths = [
        logo_path,
        f"/app/storage/logos/{logo_path}" if logo_path and not logo_path.startswith('/') else None,
        str(Path(__file__).parent.parent.parent.parent.parent / "storage" / "logos" / "niederwieser_logo.png"),
        str(Path(__file__).parent.parent.parent.parent.parent / "storage" / "logos" / f"{country_code.lower()}" / "logo.png"),
    ]

    for lp in possible_logo_paths:
        if lp and os.path.exists(lp):
            try:
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = header_para.add_run()
                # Professional logo size from settings (default 5 cm)
                logo_width = design_settings.get('logo_width_cm', 5.0)
                run.add_picture(lp, width=Cm(logo_width))

                # Add spacing line after logo for visual separation from content
                spacer_para = header.add_paragraph()
                spacer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                logo_added = True
                logger.info(f"Logo added from: {lp}")
                break
            except Exception as e:
                logger.warning(f"Could not add logo from {lp}: {e}")

    if not logo_added:
        logger.warning(f"No logo found for country {country_code}. Checked paths: {possible_logo_paths}")

    # Add document title - centered, larger font (LOCALIZED)
    title = doc.add_paragraph()
    doc_title = design_settings.get('document_title', loc_text["title"])
    title_run = title.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(16)  # Larger title
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)  # Space after title

    # "Zwischen/Tra" - centered (LOCALIZED)
    between_para = doc.add_paragraph()
    between_run = between_para.add_run(loc_text["between"])
    between_run.font.size = Pt(11)
    between_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    between_para.paragraph_format.space_after = Pt(12)

    # Employer info - centered with company name bold
    company_para = doc.add_paragraph()
    company_name = design_settings.get('company_name', 'Niederwieser GmbH' if country_code == "DE" else 'NIEDERWIESER SPA')
    address_line1 = design_settings.get('header_line1', 'Gewerbepark 9' if country_code == "DE" else 'Via Vurza 12')
    company_run = company_para.add_run(f"{company_name}\n{address_line1}")
    company_run.bold = True
    company_run.font.size = Pt(11)
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # City and postal code
    city_para = doc.add_paragraph()
    address_line2 = design_settings.get('header_line2', '87477 Sulzberg' if country_code == "DE" else '39055 Laives (BZ)')
    city_run = city_para.add_run(address_line2)
    city_run.font.size = Pt(11)
    city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Employer label - italic for legal style (LOCALIZED)
    ag_label = doc.add_paragraph()
    ag_run = ag_label.add_run(loc_text["employer_label"])
    ag_run.font.size = Pt(10)
    ag_run.italic = True
    ag_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ag_label.paragraph_format.space_after = Pt(18)

    # "und/e" centered (LOCALIZED)
    and_para = doc.add_paragraph()
    and_run = and_para.add_run(loc_text["and"])
    and_run.font.size = Pt(11)
    and_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    and_para.paragraph_format.space_after = Pt(12)

    # Employee info - centered with name bold
    employee_name = f"{form_data.get('vorname', '')} {form_data.get('nachname', '')}".strip()
    employee_street = (
        form_data.get('mitarbeiter_adresse') or
        form_data.get('adresse') or
        form_data.get('strasse', '')
    )
    employee_plz_ort = form_data.get('plz_ort', '')

    emp_para = doc.add_paragraph()
    emp_prefix = loc_text["employee_prefix"]
    if employee_street:
        emp_run = emp_para.add_run(f"{emp_prefix} {employee_name}\n{employee_street}")
    else:
        emp_run = emp_para.add_run(f"{emp_prefix} {employee_name}")
    emp_run.bold = True
    emp_run.font.size = Pt(11)
    emp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Employee city (if provided)
    if employee_plz_ort:
        emp_city_para = doc.add_paragraph()
        emp_city_run = emp_city_para.add_run(employee_plz_ort)
        emp_city_run.font.size = Pt(11)
        emp_city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Employee label - italic for legal style (LOCALIZED)
    an_label = doc.add_paragraph()
    an_run = an_label.add_run(loc_text["employee_label"])
    an_run.font.size = Pt(10)
    an_run.italic = True
    an_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    an_label.paragraph_format.space_after = Pt(24)  # Space before clauses

    # Add clauses - using enhanced HTML converter with table and nested list support
    html_converter = HtmlToDocxConverter(
        font_size_pt=design_settings.get('font_size_pt', 11),
        font_family=design_settings.get('font_family', 'Arial')
    )

    for clause in clauses:
        content = clause.get('content', '')
        # Add clause content with full HTML support (tables, nested lists, etc.)
        html_converter.convert(doc, content, form_data, country_code)
        # Spacing between clauses
        doc.add_paragraph()

    # Add signature section
    doc.add_paragraph()
    doc.add_paragraph()

    # Create signature table with 3 rows x 2 columns
    sig_table = doc.add_table(rows=3, cols=2)

    # Row 0: Date location (LOCALIZED)
    sig_location = loc_text["signature_location"]
    sig_table.cell(0, 0).text = f"{sig_location}, {loc_text['date_placeholder']}"
    sig_table.cell(0, 1).text = ""

    # Row 1: Empty for spacing
    sig_table.cell(1, 0).text = ""
    sig_table.cell(1, 1).text = ""

    # Row 2: Signature lines with names
    signatory_name = design_settings.get('signatory_name', form_data.get('signatory_name', 'Geschäftsführer'))
    company_name = design_settings.get('company_name', 'Niederwieser GmbH' if country_code == "DE" else 'NIEDERWIESER SPA')
    employee_name = f"{form_data.get('vorname', '')} {form_data.get('nachname', '')}".strip()

    sig_table.cell(2, 0).text = f"_________________\n{signatory_name}\n{company_name}"
    sig_table.cell(2, 1).text = f"_________________\n{employee_name}"

    # Save document
    doc.save(output_path)
    logger.info(f"Document created at: {output_path}")
    return output_path


def html_to_docx_paragraph_justified(doc: Document, html_content: str, form_data: dict, country_code: str = "DE"):
    """
    Convert HTML content to DOCX paragraphs.
    § titles (DE) or Art. titles (IT) are centered and bold, content is justified.
    Professional 11pt font size for readability.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_ALIGN

    # First, render placeholders with extended formatting
    rendered_html = render_placeholders_extended(html_content, form_data, country_code)

    # Parse HTML
    soup = BeautifulSoup(rendered_html, 'html.parser')

    # If no block-level elements, process the whole text
    block_elements = soup.find_all(['p', 'div'], recursive=False)
    if not block_elements:
        # Treat the whole content as a single block
        block_elements = [soup]

    # Get text and split by <br> to create separate paragraphs
    for element in block_elements:
        # Convert <br> to newlines for splitting
        for br in element.find_all('br'):
            br.replace_with('\n')

        # Get the full text
        full_text = element.get_text()

        # Split into lines (each <br> becomes a separate paragraph)
        lines = full_text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            para = doc.add_paragraph()

            # Check if this is a title (§ for DE, Art. for IT, or all-caps)
            is_title = (
                line.startswith('§') or
                line.upper().startswith('ART.') or
                line.upper().startswith('ARTICOLO') or
                (line.isupper() and len(line) < 80)  # All-caps short lines are likely titles
            )

            if is_title:
                para.alignment = WD_ALIGN.CENTER
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
            else:
                # Regular content - justified, 11pt for readability
                para.alignment = WD_ALIGN.JUSTIFY
                run = para.add_run(line)
                run.font.size = Pt(11)


@router.post("/generate/{document_type_id}")
async def generate_document_by_type(
    document_type_id: int,
    request_data: GenerateDocumentRequest,
    background_tasks: BackgroundTasks,
    current_user: Annotated[models.User, Depends(deps.get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """
    Generate a document from a document type ID and context data.

    This endpoint assembles clauses from the database (like the preview does)
    and generates a DOCX/PDF document with placeholders replaced.
    """
    # 1. Load document type from DB
    doc_type = await db.get(DocumentType, document_type_id)
    if not doc_type or not doc_type.is_active:
        logger.warning(
            f"Invalid document type requested: {document_type_id} by user {current_user.id}"
        )
        raise HTTPException(
            status_code=404,
            detail="Dokumenttyp nicht gefunden oder inaktiv"
        )

    country_code = doc_type.country_code or "DE"

    # 2. Load design settings
    design_result = await db.execute(
        select(models.DesignSetting).where(models.DesignSetting.country_code == country_code)
    )
    design = design_result.scalar_one_or_none()

    design_settings = {
        # Company info
        "company_name": design.company_name if design else "Niederwieser GmbH",
        "header_line1": design.header_line1 if design else "Gewerbepark 9",
        "header_line2": design.header_line2 if design else "87477 Sulzberg",
        "font_family": design.font_family if design else "Arial",
        "document_title": doc_type.name or "Arbeitsvertrag",
        "logo_path": design.logo_path if design else None,

        # Signatory from request (required field)
        "signatory_name": request_data.signatory_name,

        # DIN 5008 Page Format Settings
        "margin_left_cm": float(getattr(design, 'margin_left_cm', None) or 2.5),
        "margin_right_cm": float(getattr(design, 'margin_right_cm', None) or 2.0),
        "margin_top_cm": float(getattr(design, 'margin_top_cm', None) or 2.5),
        "margin_bottom_cm": float(getattr(design, 'margin_bottom_cm', None) or 2.0),
        "header_distance_cm": float(getattr(design, 'header_distance_cm', None) or 1.25),
        "footer_distance_cm": float(getattr(design, 'footer_distance_cm', None) or 1.0),

        # Typography
        "font_size_pt": int(getattr(design, 'font_size_pt', None) or 11),
        "line_spacing": float(getattr(design, 'line_spacing', None) or 1.15),

        # Logo
        "logo_width_cm": float(getattr(design, 'logo_width_cm', None) or 5.0),
        "logo_position": getattr(design, 'logo_position', None) or "right",
    }

    # 3. Load clauses for this document type
    clause_refs_result = await db.execute(
        select(DocumentTypeClause)
        .where(DocumentTypeClause.document_type_id == document_type_id)
        .order_by(DocumentTypeClause.display_order)
    )
    clause_refs = clause_refs_result.scalars().all()

    # 4. Build form_data context from request
    form_data = request_data.model_dump(
        exclude={"output_format", "attachment_ids", "clause_ids"},
        exclude_none=True
    )

    # 5. Filter and assemble clauses
    active_clauses = []
    for ref in clause_refs:
        clause = await db.get(Clause, ref.clause_id)
        if clause and clause.is_active:
            # Parse condition if exists
            condition = None
            if ref.condition:
                try:
                    condition = json.loads(ref.condition) if isinstance(ref.condition, str) else ref.condition
                except (json.JSONDecodeError, TypeError):
                    condition = None

            # Evaluate condition
            if evaluate_condition(condition, form_data):
                active_clauses.append({
                    "id": clause.id,
                    "title": clause.title,
                    "content": clause.content_html,
                    "is_mandatory": ref.is_mandatory,
                })

    # 6. Generate document
    try:
        nachname = form_data.get('nachname', 'Dokument')
        # Sanitize filename
        safe_nachname = re.sub(r'[^\w\s-]', '', nachname).strip().replace(' ', '_')
        output_filename = f"Vertrag_{safe_nachname}.docx"
        output_path = OUTPUT_DIR / output_filename

        # Create DOCX from clauses
        docx_path = create_document_from_clauses(
            clauses=active_clauses,
            form_data=form_data,
            country_code=country_code,
            design_settings=design_settings,
            output_path=output_path
        )

        # Convert to PDF if requested
        if request_data.output_format == "pdf":
            if request_data.async_pdf:
                # Async PDF conversion via Celery
                from app.tasks.pdf_tasks import convert_docx_to_pdf, update_pdf_job_status
                import uuid

                job_id = str(uuid.uuid4())
                update_pdf_job_status(job_id, "pending")

                # Queue the conversion task
                convert_docx_to_pdf.delay(
                    docx_path=str(docx_path),
                    job_id=job_id,
                    cleanup_docx=True
                )

                return AsyncJobResponse(
                    job_id=job_id,
                    status="pending",
                    message="PDF conversion started. Poll /api/v1/generate/jobs/{job_id} for status."
                )
            else:
                # Synchronous conversion (blocking - use for small documents only)
                output_path = engine.convert_to_pdf(docx_path)
                media_type = "application/pdf"
                output_filename = output_filename.replace('.docx', '.pdf')
        else:
            output_path = docx_path
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return FileResponse(
            output_path,
            media_type=media_type,
            filename=output_filename
        )

    except Exception as e:
        logger.exception(f"Document generation failed for user {current_user.id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Dokumentgenerierung fehlgeschlagen: {str(e)}"
        )


@router.get("/jobs/{job_id}", response_model=JobStatusResponse)
async def get_pdf_job_status(
    job_id: str,
    current_user: Annotated[models.User, Depends(deps.get_current_user)],
):
    """
    Get the status of an async PDF conversion job.

    Poll this endpoint after submitting a document generation request with async_pdf=true.

    Returns:
        - status: pending, processing, completed, failed
        - output_path: Path to the generated PDF (when completed)
        - download_url: URL to download the PDF (when completed)
        - error: Error message (when failed)
    """
    from app.tasks.pdf_tasks import get_pdf_job_status as get_status

    job_data = get_status(job_id)

    response = JobStatusResponse(
        job_id=job_id,
        status=job_data.get("status", "unknown"),
        output_path=job_data.get("output_path"),
        error=job_data.get("error"),
    )

    # Add download URL if completed
    if response.status == "completed" and response.output_path:
        response.download_url = f"/api/v1/generate/download/{job_id}"

    return response


@router.get("/download/{job_id}")
async def download_pdf_by_job_id(
    job_id: str,
    current_user: Annotated[models.User, Depends(deps.get_current_user)],
):
    """
    Download a PDF file from a completed async conversion job.

    Use this after the job status shows 'completed'.
    """
    from app.tasks.pdf_tasks import get_pdf_job_status as get_status

    job_data = get_status(job_id)

    if job_data.get("status") != "completed":
        raise HTTPException(
            status_code=400,
            detail=f"Job not completed. Current status: {job_data.get('status', 'unknown')}"
        )

    output_path = job_data.get("output_path")
    if not output_path or not Path(output_path).exists():
        raise HTTPException(
            status_code=404,
            detail="PDF file not found. It may have been cleaned up."
        )

    return FileResponse(
        output_path,
        media_type="application/pdf",
        filename=Path(output_path).name
    )


# Legacy endpoint für Abwärtskompatibilität
@router.post("/generate/legacy/{template_name}")
async def generate_document_legacy(
    template_name: str,
    context: dict,
    background_tasks: BackgroundTasks,
    current_user: Annotated[models.User, Depends(deps.get_current_user)],
):
    """
    Legacy endpoint: Generate a PDF document from a template name.

    DEPRECATED: Use /generate/{document_type_id} instead.
    """
    # SECURITY: Template-Name gegen Whitelist validieren
    ALLOWED_TEMPLATES = {
        "arbeitsvertrag", "kuendigung", "zeugnis", "vertrag",
        "contract", "letter", "certificate", "amendment", "template", "default",
    }

    if not re.match(r'^[a-zA-Z0-9_äöüÄÖÜß-]+$', template_name):
        raise HTTPException(status_code=400, detail="Ungültiger Template-Name")

    base_name = template_name.lower().replace('.docx', '').replace('.pdf', '')
    if base_name not in ALLOWED_TEMPLATES:
        logger.warning(f"Invalid template requested: {template_name} by user {current_user.id}")
        raise HTTPException(status_code=400, detail="Ungültiger Template-Name")

    try:
        output_filename = f"gen_{current_user.id}_{template_name}"
        docx_path = engine.assemble_document(template_name, context, output_filename + ".docx")
        pdf_path = engine.convert_to_pdf(docx_path)

        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=f"{output_filename}.pdf"
        )
    except FileNotFoundError:
        logger.error(f"Template not found: {template_name}")
        raise HTTPException(status_code=404, detail="Template nicht gefunden")
    except Exception as e:
        logger.exception(f"Document generation failed for user {current_user.id}: {e}")
        raise HTTPException(
            status_code=500,
            detail="Dokumentgenerierung fehlgeschlagen. Bitte versuchen Sie es erneut."
        )
