"""
User Templates API: CRUD operations for user-uploaded DOCX templates.

Users can upload their own branded DOCX templates (with logo, header/footer)
to use as layout basis when generating documents.

Templates can be scoped as:
- 'company': visible to all users (default)
- 'team': visible to members of a specific team
- 'private': visible only to the creator

Template types:
- 'stationery': Blanko-Briefpapier (nur Layout, Header/Footer)
- 'content': Inhaltsvorlage (enthält bereits Text-Bausteine)
"""
import os
import re
import uuid
import logging
import aiofiles
from datetime import datetime
from pathlib import Path
from typing import Annotated, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete, or_, and_

from app.db import get_db
from app.models.core import User
from app.models.user_templates import UserTemplate
from app.models.enterprise import TeamMember
from app.api.deps import get_current_user

router = APIRouter()
logger = logging.getLogger(__name__)

# Storage path for user templates
BASE_DIR = Path(__file__).parent.parent.parent.parent.parent
USER_TEMPLATE_STORAGE = BASE_DIR / "storage" / "user-templates"
USER_TEMPLATE_STORAGE.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

# DOCX magic bytes (PK zip header)
DOCX_MAGIC = b"PK\x03\x04"


async def _can_access_template(db: AsyncSession, template: UserTemplate, user: User) -> bool:
    """Check if user can access this template based on scope."""
    # Owner always has access
    if template.user_id == user.id:
        return True
    # Company scope: everyone
    if template.scope == "company":
        return True
    # Team scope: check membership
    if template.scope == "team" and template.team_id:
        user_id_str = str(user.id)
        result = await db.execute(
            select(TeamMember.id).where(
                and_(
                    TeamMember.team_id == template.team_id,
                    TeamMember.user_id == user_id_str,
                )
            ).limit(1)
        )
        return result.scalar_one_or_none() is not None
    # Private: only owner (handled above)
    return False


def _template_to_dict(t: UserTemplate, current_user_id: int) -> dict:
    """Convert a UserTemplate to a response dict."""
    # Thumbnail-URL konstruieren (falls vorhanden)
    thumbnail_url = None
    if t.thumbnail_path:
        thumbnail_url = f"/api/v1/user-templates/{t.id}/thumbnail"

    return {
        "id": t.id,
        "name": t.name,
        "description": t.description,
        "original_filename": t.original_filename,
        "file_size": t.file_size,
        "country_code": t.country_code,
        "has_header": t.has_header,
        "has_footer": t.has_footer,
        "has_logo": t.has_logo,
        "font_family": t.font_family,
        "scope": t.scope,
        "team_id": t.team_id,
        "category": t.category,
        "template_type": t.template_type,
        "is_default": t.is_default,
        "thumbnail_url": thumbnail_url,
        "is_own": t.user_id == current_user_id,
        "created_at": t.created_at.isoformat() if t.created_at else None,
        "updated_at": t.updated_at.isoformat() if t.updated_at else None,
    }


def _analyze_docx(file_path: Path) -> dict:
    """Analyze a DOCX file to extract metadata about headers, footers, logos."""
    try:
        from docx import Document
        doc = Document(str(file_path))

        has_header = False
        has_footer = False
        has_logo = False
        font_family = None

        # Check sections for headers/footers
        for section in doc.sections:
            header = section.header
            if header and header.paragraphs:
                for para in header.paragraphs:
                    if para.text.strip():
                        has_header = True
                    for run in para.runs:
                        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            has_logo = True
                            has_header = True

            footer = section.footer
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    if para.text.strip():
                        has_footer = True
                    for run in para.runs:
                        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            has_footer = True

        # Check images in the document package (rels)
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            rels = doc.part.rels
            for rel in rels.values():
                if "image" in rel.reltype:
                    has_logo = True
                    break
        except Exception:
            pass

        # Detect font from Normal style
        try:
            style = doc.styles['Normal']
            if style.font and style.font.name:
                font_family = style.font.name
        except Exception:
            pass

        # Fallback: check first paragraph run font
        if not font_family:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font and run.font.name:
                        font_family = run.font.name
                        break
                if font_family:
                    break

        return {
            "has_header": has_header,
            "has_footer": has_footer,
            "has_logo": has_logo,
            "font_family": font_family,
        }
    except Exception as e:
        logger.warning(f"Could not analyze DOCX: {e}")
        return {
            "has_header": False,
            "has_footer": False,
            "has_logo": False,
            "font_family": None,
        }


def _generate_thumbnail_background(template_id: int, docx_path: Path, db_url: str):
    """
    Background-Task: Thumbnail generieren und Pfad in DB speichern.

    Wird nach dem Upload asynchron ausgeführt, damit der Upload-Request
    nicht auf die LibreOffice-Konvertierung warten muss.
    """
    from app.services.thumbnail_service import generate_thumbnail

    thumb_path = generate_thumbnail(docx_path)
    if thumb_path is None:
        logger.warning(f"Thumbnail-Generierung für Template {template_id} fehlgeschlagen")
        return

    # Thumbnail-Pfad in DB speichern (synchron, da BackgroundTask)
    try:
        import asyncio
        from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession as AS
        from sqlalchemy.orm import sessionmaker

        async def _save_path():
            engine = create_async_engine(db_url, pool_size=1, max_overflow=0)
            async_session = sessionmaker(engine, class_=AS, expire_on_commit=False)
            try:
                async with async_session() as session:
                    template = await session.get(UserTemplate, template_id)
                    if template:
                        template.thumbnail_path = str(thumb_path)
                        await session.commit()
                        logger.info(f"Thumbnail-Pfad gespeichert für Template {template_id}")
                    else:
                        # Template wurde zwischenzeitlich gelöscht — Thumbnail-Datei aufräumen
                        if thumb_path.exists():
                            thumb_path.unlink(missing_ok=True)
                            logger.info(f"Verwaistes Thumbnail gelöscht für Template {template_id}")
            finally:
                await engine.dispose()

        asyncio.run(_save_path())
    except Exception as e:
        logger.error(f"Fehler beim Speichern des Thumbnail-Pfads: {e}")


# ═══════════════════════════════════════════════════════════════════════════
# FIXED-PATH ENDPOINTS (müssen VOR /{template_id} registriert werden)
# ═══════════════════════════════════════════════════════════════════════════

@router.get("/suggest")
async def suggest_template(
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
    country_code: Optional[str] = None,
    category: Optional[str] = None,
):
    """
    Schlägt die passende Standard-Briefpapier-Vorlage vor.

    Logik:
    1. Suche nach is_default=True für das Land/die Kategorie
    2. Fallback: neueste Briefpapier-Vorlage für das Land
    """
    user_id_str = str(current_user.id)

    # Zugriffsfilter (wie bei list)
    team_ids_result = await db.execute(
        select(TeamMember.team_id).where(TeamMember.user_id == user_id_str)
    )
    user_team_ids = [row[0] for row in team_ids_result.all()]

    access_conditions = [
        UserTemplate.user_id == current_user.id,
        UserTemplate.scope == "company",
    ]
    if user_team_ids:
        access_conditions.append(
            and_(
                UserTemplate.scope == "team",
                UserTemplate.team_id.in_(user_team_ids),
            )
        )

    # Nur Briefpapier-Vorlagen
    base_filter = and_(
        or_(*access_conditions),
        UserTemplate.template_type == "stationery",
    )

    # 1. Versuch: Standard-Vorlage für Land/Kategorie
    query = select(UserTemplate).where(
        and_(base_filter, UserTemplate.is_default == True)
    )
    if country_code:
        query = query.where(UserTemplate.country_code == country_code.upper())
    if category:
        query = query.where(UserTemplate.category == category)

    result = await db.execute(query.limit(1))
    template = result.scalar_one_or_none()

    # 2. Fallback: neueste Briefpapier-Vorlage für das Land
    if not template:
        query = select(UserTemplate).where(base_filter)
        if country_code:
            query = query.where(UserTemplate.country_code == country_code.upper())
        query = query.order_by(UserTemplate.created_at.desc())

        result = await db.execute(query.limit(1))
        template = result.scalar_one_or_none()

    if not template:
        return {"template": None, "message": "Keine passende Briefpapier-Vorlage gefunden"}

    return {"template": _template_to_dict(template, current_user.id)}


# ═══════════════════════════════════════════════════════════════════════════
# LIST & CREATE
# ═══════════════════════════════════════════════════════════════════════════

@router.get("")
async def list_user_templates(
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
    country_code: Optional[str] = None,
    scope: Optional[str] = None,
    team_id: Optional[int] = None,
    category: Optional[str] = None,
    template_type: Optional[str] = None,
):
    """List all templates accessible to the current user."""
    user_id_str = str(current_user.id)

    # Get user's team IDs (TeamMember.user_id is String(255))
    team_ids_result = await db.execute(
        select(TeamMember.team_id).where(TeamMember.user_id == user_id_str)
    )
    user_team_ids = [row[0] for row in team_ids_result.all()]

    # Build access filter: own OR company OR team-accessible
    access_conditions = [
        UserTemplate.user_id == current_user.id,
        UserTemplate.scope == "company",
    ]
    if user_team_ids:
        access_conditions.append(
            and_(
                UserTemplate.scope == "team",
                UserTemplate.team_id.in_(user_team_ids),
            )
        )

    query = select(UserTemplate).where(
        or_(*access_conditions)
    ).order_by(UserTemplate.created_at.desc())

    # Apply optional filters
    if country_code:
        query = query.where(UserTemplate.country_code == country_code.upper())
    if scope:
        query = query.where(UserTemplate.scope == scope)
    if team_id:
        query = query.where(UserTemplate.team_id == team_id)
    if category:
        query = query.where(UserTemplate.category == category)
    if template_type:
        query = query.where(UserTemplate.template_type == template_type)

    result = await db.execute(query)
    templates = result.scalars().all()

    return {
        "items": [_template_to_dict(t, current_user.id) for t in templates],
        "total": len(templates),
    }


@router.post("")
async def upload_user_template(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    name: str = Query(..., min_length=1, max_length=255, description="Name der Vorlage"),
    description: Optional[str] = Query(None, max_length=1000),
    country_code: Optional[str] = Query(None, max_length=2),
    scope: str = Query("company", pattern="^(company|team|private)$"),
    team_id: Optional[int] = Query(None),
    category: Optional[str] = Query(None, max_length=100),
    template_type: str = Query("stationery", pattern="^(stationery|content)$"),
    is_default: bool = Query(False, description="Als Standard-Briefpapier setzen"),
    current_user: Annotated[User, Depends(get_current_user)] = None,
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a new DOCX template.

    The template will be analyzed for headers, footers, logos and font.
    For stationery templates, a thumbnail is generated in the background.
    """
    # Validate file type
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Nur DOCX-Dateien sind erlaubt")

    # Read content
    content = await file.read()

    # Validate size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Datei zu gross. Maximum: 10 MB")

    # Validate magic bytes (DOCX is a ZIP file)
    if not content[:4] == DOCX_MAGIC:
        raise HTTPException(status_code=400, detail="Ungültige DOCX-Datei")

    # Validate country code if provided
    if country_code:
        country_code = country_code.upper()
        if not re.match(r'^[A-Z]{2}$', country_code):
            raise HTTPException(status_code=400, detail="Ungültiger Ländercode (2 Buchstaben)")

    # Validate scope + team_id combination
    if scope == "team" and not team_id:
        raise HTTPException(status_code=400, detail="team_id ist erforderlich für Team-Vorlagen")
    if scope != "team" and team_id:
        raise HTTPException(status_code=400, detail="team_id nur für Team-Vorlagen erlaubt")

    # Verify team membership if team scope
    if team_id:
        user_id_str = str(current_user.id)
        role_result = await db.execute(
            select(TeamMember.role).where(
                and_(
                    TeamMember.team_id == team_id,
                    TeamMember.user_id == user_id_str,
                )
            ).limit(1)
        )
        role = role_result.scalar_one_or_none()
        if not role:
            raise HTTPException(status_code=403, detail="Kein Zugriff auf dieses Team")

    # Create user directory
    user_dir = USER_TEMPLATE_STORAGE / str(current_user.id)
    user_dir.mkdir(parents=True, exist_ok=True)

    # Generate unique filename
    file_uuid = str(uuid.uuid4())
    stored_filename = f"{file_uuid}.docx"
    file_path = user_dir / stored_filename

    # Verify path is within storage directory (prevent traversal)
    if not file_path.resolve().is_relative_to(USER_TEMPLATE_STORAGE.resolve()):
        raise HTTPException(status_code=400, detail="Ungültiger Dateipfad")

    # Save file
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)

    # Analyze DOCX structure
    metadata = _analyze_docx(file_path)

    # Create database record
    template = UserTemplate(
        user_id=current_user.id,
        name=name,
        description=description,
        original_filename=file.filename,
        stored_filename=stored_filename,
        file_size=len(content),
        country_code=country_code,
        scope=scope,
        team_id=team_id if scope == "team" else None,
        category=category,
        template_type=template_type,
        has_header=metadata["has_header"],
        has_footer=metadata["has_footer"],
        has_logo=metadata["has_logo"],
        font_family=metadata["font_family"],
    )
    db.add(template)
    await db.commit()
    await db.refresh(template)

    # Falls is_default gesetzt: andere Defaults für gleiches Land zurücksetzen
    if is_default and template_type == "stationery":
        from sqlalchemy import update
        country_filter = (
            UserTemplate.country_code == country_code
            if country_code is not None
            else UserTemplate.country_code.is_(None)
        )
        await db.execute(
            update(UserTemplate)
            .where(
                and_(
                    country_filter,
                    UserTemplate.is_default == True,
                    UserTemplate.id != template.id,
                    UserTemplate.template_type == "stationery",
                )
            )
            .values(is_default=False)
        )
        template.is_default = True
        await db.commit()
        await db.refresh(template)

    logger.info(
        f"User template uploaded: id={template.id}, user={current_user.id}, "
        f"name={name}, scope={scope}, type={template_type}"
    )

    # Thumbnail im Hintergrund generieren (nur für Briefpapier)
    if template_type == "stationery":
        from app.core.config import settings
        background_tasks.add_task(
            _generate_thumbnail_background,
            template.id,
            file_path,
            settings.DATABASE_URL,
        )

    return _template_to_dict(template, current_user.id)


# ═══════════════════════════════════════════════════════════════════════════
# SINGLE TEMPLATE ENDPOINTS (/{template_id}/...)
# ═══════════════════════════════════════════════════════════════════════════

@router.get("/{template_id}")
async def get_user_template(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """Get metadata for a single user template."""
    template = await db.get(UserTemplate, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    if not await _can_access_template(db, template, current_user):
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    return _template_to_dict(template, current_user.id)


@router.get("/{template_id}/download")
async def download_user_template(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """Download a user template file."""
    template = await db.get(UserTemplate, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    if not await _can_access_template(db, template, current_user):
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    # Use uploader's user_id for file path (not current_user)
    file_path = USER_TEMPLATE_STORAGE / str(template.user_id) / template.stored_filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template-Datei nicht gefunden")

    return FileResponse(
        str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=template.original_filename,
    )


@router.get("/{template_id}/thumbnail")
async def get_template_thumbnail(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """
    Gibt das PNG-Thumbnail einer Briefpapier-Vorlage zurück.

    Falls kein Thumbnail existiert, wird versucht es on-the-fly zu generieren.
    """
    template = await db.get(UserTemplate, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    if not await _can_access_template(db, template, current_user):
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    # Thumbnail-Pfad prüfen (mit Path-Traversal-Schutz)
    thumb_path = None
    if template.thumbnail_path:
        thumb_path = Path(template.thumbnail_path).resolve()
        # Sicherheitscheck: Pfad muss innerhalb des Storage-Verzeichnisses liegen
        from app.services.thumbnail_service import THUMBNAIL_STORAGE
        if not thumb_path.is_relative_to(THUMBNAIL_STORAGE.resolve()):
            logger.warning(f"Path traversal attempt: {template.thumbnail_path}")
            thumb_path = None

    # Falls kein Thumbnail vorhanden, on-the-fly generieren
    if not thumb_path or not thumb_path.exists():
        from app.services.thumbnail_service import generate_thumbnail

        docx_path = USER_TEMPLATE_STORAGE / str(template.user_id) / template.stored_filename
        if not docx_path.exists():
            raise HTTPException(status_code=404, detail="Template-Datei nicht gefunden")

        thumb_path = generate_thumbnail(docx_path)
        if thumb_path is None:
            raise HTTPException(
                status_code=503,
                detail="Thumbnail-Generierung fehlgeschlagen (LibreOffice nicht verfügbar?)"
            )

        # Pfad in DB speichern für nächstes Mal
        template.thumbnail_path = str(thumb_path)
        await db.commit()

    return FileResponse(
        str(thumb_path),
        media_type="image/png",
        filename=f"thumbnail_{template_id}.png",
    )


@router.get("/{template_id}/zones")
async def get_template_zones(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """
    Extrahiert Header-/Footer-Bilder, Text und Seitenränder aus einer DOCX-Vorlage.

    Wird für die Canvas-Vorschau im Frontend verwendet.
    """
    template = await db.get(UserTemplate, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    if not await _can_access_template(db, template, current_user):
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    file_path = USER_TEMPLATE_STORAGE / str(template.user_id) / template.stored_filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template-Datei nicht gefunden")

    from app.services.user_template_service import extract_template_zones
    try:
        zones = extract_template_zones(file_path)
    except Exception as e:
        logger.error(f"Fehler beim Extrahieren der Zonen aus Template {template_id}: {e}")
        raise HTTPException(
            status_code=422,
            detail="Die DOCX-Datei konnte nicht verarbeitet werden. Möglicherweise ist sie beschädigt."
        )

    return {
        "template_id": template.id,
        "template_name": template.name,
        **zones,
    }


@router.put("/{template_id}/default")
async def set_default_template(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """
    Markiert eine Briefpapier-Vorlage als Standard für ihr Land.

    Entfernt is_default von allen anderen Vorlagen mit gleichem country_code.
    Nur der Ersteller oder Nutzer mit company-scope dürfen das setzen.
    """
    template = await db.get(UserTemplate, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    # Nur Owner oder company-scope Vorlagen können als Default gesetzt werden
    is_owner = template.user_id == current_user.id
    is_company_scope = template.scope == "company"

    if not is_owner and not is_company_scope:
        raise HTTPException(
            status_code=403,
            detail="Nur der Ersteller oder Admins können Standardvorlagen festlegen"
        )

    # Alle anderen Vorlagen mit gleichem country_code auf is_default=False setzen
    country_filter = (
        UserTemplate.country_code == template.country_code
        if template.country_code is not None
        else UserTemplate.country_code.is_(None)
    )
    # Atomares Update: Alle anderen Defaults in einem Statement zurücksetzen
    from sqlalchemy import update
    await db.execute(
        update(UserTemplate)
        .where(
            and_(
                country_filter,
                UserTemplate.is_default == True,
                UserTemplate.id != template_id,
                UserTemplate.template_type == "stationery",
            )
        )
        .values(is_default=False)
    )

    # Diese Vorlage als Standard setzen
    template.is_default = True
    await db.commit()
    await db.refresh(template)

    logger.info(
        f"Template {template_id} als Standard gesetzt "
        f"(country={template.country_code}, user={current_user.id})"
    )

    return _template_to_dict(template, current_user.id)


@router.patch("/{template_id}")
async def update_user_template(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
    name: Optional[str] = Query(None, min_length=1, max_length=255),
    description: Optional[str] = Query(None, max_length=1000),
    country_code: Optional[str] = Query(None, max_length=2),
    category: Optional[str] = Query(None, max_length=100),
):
    """
    Aktualisiert Name, Beschreibung, Land oder Kategorie einer Vorlage.

    Nur der Ersteller kann seine Vorlagen bearbeiten.
    """
    template = await db.get(UserTemplate, template_id)
    if not template or template.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    if name is not None:
        template.name = name.strip()
    if description is not None:
        template.description = description.strip() or None
    if country_code is not None:
        cc = country_code.strip().upper()
        if cc and not re.match(r'^[A-Z]{2}$', cc):
            raise HTTPException(status_code=400, detail="Ungültiger Ländercode (2 Buchstaben)")
        template.country_code = cc or None
    if category is not None:
        template.category = category.strip() or None

    await db.commit()
    await db.refresh(template)

    logger.info(f"User template updated: id={template_id}, user={current_user.id}")

    return _template_to_dict(template, current_user.id)


@router.delete("/{template_id}")
async def delete_user_template(
    template_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: AsyncSession = Depends(get_db),
):
    """Delete a user template. Only the creator can delete."""
    template = await db.get(UserTemplate, template_id)
    if not template or template.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Vorlage nicht gefunden")

    # Remove file
    file_path = USER_TEMPLATE_STORAGE / str(current_user.id) / template.stored_filename
    if file_path.exists():
        os.remove(file_path)

    # Remove thumbnail if exists
    if template.thumbnail_path:
        thumb_path = Path(template.thumbnail_path)
        if thumb_path.exists():
            os.remove(thumb_path)

    # Remove DB record
    await db.delete(template)
    await db.commit()

    logger.info(f"User template deleted: id={template_id}, user={current_user.id}")

    return {"status": "deleted", "id": template_id}
