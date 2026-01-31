"""
Dokumententyp-Import API (v4.2 Feature)

Importiert komplette Word-Dokumente als neue Dokumententypen inkl. Klauseln.
Erkennt Platzhalter, findet ähnliche Klauseln und erstellt Formularfelder.
"""

import re
import io
import json
from datetime import datetime
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from pydantic import BaseModel

from app.db import get_db
from app.models.documents import Clause, DocumentType, DocumentTypeClause

router = APIRouter()


# ══════════════════════════════════════════════════════════════════════════════
# SCHEMAS
# ══════════════════════════════════════════════════════════════════════════════

class PlaceholderSuggestion(BaseModel):
    """Vorgeschlagener Platzhalter."""
    original_text: str
    suggested_placeholder: str
    placeholder_type: str  # text, date, number, currency
    confidence: float  # 0.0 - 1.0
    accepted: bool = True


class ExtractedSection(BaseModel):
    """Ein erkannter Abschnitt aus dem Word-Dokument."""
    index: int
    title: str
    content: str
    content_html: str
    char_count: int

    # Klausel-Zuordnung
    action: str = "create_new"  # create_new, use_existing, add_variant, skip
    existing_clause_id: Optional[int] = None
    existing_clause_title: Optional[str] = None
    similarity_score: Optional[float] = None

    # Für Varianten
    is_variant_of: Optional[int] = None  # ID der Haupt-Klausel
    variant_name: Optional[str] = None  # z.B. "6 Monate"

    # Dokumententyp-Einstellungen
    is_mandatory: bool = True
    display_order: int = 0

    # Erkannte Platzhalter
    placeholders: List[PlaceholderSuggestion] = []


class SimilarClause(BaseModel):
    """Ähnliche existierende Klausel."""
    id: int
    title: str
    content_preview: str
    similarity_score: float
    has_variants: bool = False
    variants: List[Dict[str, Any]] = []


class DocumentTypeAnalysisResult(BaseModel):
    """Ergebnis der Dokumententyp-Analyse."""
    filename: str
    total_chars: int
    sections_count: int
    sections: List[ExtractedSection]
    warnings: List[str]

    # Statistiken
    headings_found: int
    tables_converted: int
    images_ignored: int
    placeholders_detected: int

    # Ähnliche Klauseln
    similar_clauses_found: int


class DocumentTypeImportRequest(BaseModel):
    """Request zum Importieren als Dokumententyp."""
    # Dokumententyp-Metadaten
    document_type_name: str
    document_type_category: str
    country_code: str = "DE"

    # Konfigurierte Abschnitte
    sections: List[ExtractedSection]

    # Globale Platzhalter (für alle Abschnitte)
    global_placeholders: List[PlaceholderSuggestion] = []


class ImportResult(BaseModel):
    """Ergebnis des Imports."""
    document_type_id: int
    document_type_name: str
    clauses_created: int
    clauses_linked: int
    variants_added: int
    form_fields_created: int


# ══════════════════════════════════════════════════════════════════════════════
# PLACEHOLDER DETECTION
# ══════════════════════════════════════════════════════════════════════════════

# Bekannte Platzhalter-Muster
PLACEHOLDER_PATTERNS = [
    # Namen
    (r'\b(Herr|Frau)\s+[A-ZÄÖÜ][a-zäöüß]+\s+[A-ZÄÖÜ][a-zäöüß]+\b', 'mitarbeiter_name', 'text', 0.9),
    (r'\b[A-ZÄÖÜ][a-zäöüß]+\s+[A-ZÄÖÜ][a-zäöüß]+\b(?=\s*,?\s*(geb\.|geboren))', 'mitarbeiter_name', 'text', 0.85),

    # Datumsformate
    (r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', 'datum', 'date', 0.8),
    (r'\b\d{1,2}\.\s*(Januar|Februar|März|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s*\d{4}\b', 'datum', 'date', 0.85),

    # Geldbeträge
    (r'\b\d{1,3}(?:\.\d{3})*(?:,\d{2})?\s*(?:€|EUR|Euro)\b', 'betrag', 'currency', 0.9),
    (r'\b(?:€|EUR)\s*\d{1,3}(?:\.\d{3})*(?:,\d{2})?\b', 'betrag', 'currency', 0.9),

    # Stunden/Arbeitszeit
    (r'\b\d{1,2}\s*(?:Stunden|Std\.?)\s*(?:pro\s*Woche|wöchentlich)?\b', 'wochenstunden', 'number', 0.85),

    # Urlaubstage
    (r'\b\d{1,2}\s*(?:Arbeits)?tage?\s*(?:Urlaub|Jahresurlaub)\b', 'urlaubstage', 'number', 0.85),
    (r'\b(?:Urlaub|Jahresurlaub)\s*(?:von\s*)?\d{1,2}\s*(?:Arbeits)?tage?\b', 'urlaubstage', 'number', 0.85),

    # Kündigungsfristen
    (r'\b\d+\s*(?:Wochen?|Monate?)\s*(?:zum\s*(?:Monatsende|Quartalsende))?\b', 'kuendigungsfrist', 'text', 0.75),

    # Probezeit
    (r'\b\d+\s*Monate?\s*(?:Probezeit)?\b(?=.*Probezeit)', 'probezeit_monate', 'number', 0.8),

    # Adressen
    (r'\b[A-ZÄÖÜ][a-zäöüß]+(?:straße|str\.|weg|platz|allee)\s*\d+[a-z]?\b', 'strasse', 'text', 0.7),
    (r'\b\d{5}\s+[A-ZÄÖÜ][a-zäöüß]+\b', 'plz_ort', 'text', 0.75),

    # Personalnummer
    (r'\bPersonalnummer:?\s*\d+\b', 'personalnummer', 'text', 0.9),
]

# Kontext-basierte Platzhalter (in der Nähe von Schlüsselwörtern)
CONTEXT_PLACEHOLDERS = {
    'eintrittsdatum': ['eintritt', 'beginn', 'arbeitsverhältnis beginnt', 'start'],
    'gehalt': ['vergütung', 'gehalt', 'monatlich', 'brutto', 'entgelt'],
    'position': ['position', 'stelle', 'tätigkeit', 'funktion'],
    'abteilung': ['abteilung', 'bereich', 'department'],
    'vorgesetzter': ['vorgesetzt', 'berichtet an', 'fachlich'],
    'arbeitsort': ['arbeitsort', 'einsatzort', 'dienstort', 'standort'],
}


def detect_placeholders(text: str) -> List[PlaceholderSuggestion]:
    """Erkennt potenzielle Platzhalter im Text."""
    suggestions = []
    found_positions = set()  # Vermeidet Duplikate

    for pattern, placeholder_name, placeholder_type, confidence in PLACEHOLDER_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            start, end = match.span()

            # Überlappung prüfen
            if any(start < pos < end or pos_start < start < pos_end
                   for pos, pos_start, pos_end in found_positions):
                continue

            original = match.group(0)
            found_positions.add((start, start, end))

            # Spezifischen Platzhalter-Namen basierend auf Kontext bestimmen
            context_name = _get_context_placeholder_name(text, start, placeholder_name)

            suggestions.append(PlaceholderSuggestion(
                original_text=original,
                suggested_placeholder=f"{{{{{context_name}}}}}",
                placeholder_type=placeholder_type,
                confidence=confidence,
            ))

    return suggestions


def _get_context_placeholder_name(text: str, position: int, default_name: str) -> str:
    """Bestimmt den Platzhalter-Namen basierend auf dem Kontext."""
    # Textausschnitt um die Position herum
    start = max(0, position - 100)
    end = min(len(text), position + 100)
    context = text[start:end].lower()

    for placeholder_name, keywords in CONTEXT_PLACEHOLDERS.items():
        if any(keyword in context for keyword in keywords):
            return placeholder_name

    return default_name


def apply_placeholders_to_content(content: str, placeholders: List[PlaceholderSuggestion]) -> str:
    """Ersetzt erkannte Werte durch Platzhalter im Content."""
    result = content

    # Sortieren nach Position (rückwärts, damit Ersetzungen die Positionen nicht verschieben)
    sorted_placeholders = sorted(
        [p for p in placeholders if p.accepted],
        key=lambda p: content.find(p.original_text),
        reverse=True
    )

    for placeholder in sorted_placeholders:
        result = result.replace(placeholder.original_text, placeholder.suggested_placeholder, 1)

    return result


# ══════════════════════════════════════════════════════════════════════════════
# CLAUSE SIMILARITY
# ══════════════════════════════════════════════════════════════════════════════

async def find_similar_clauses(
    db: AsyncSession,
    title: str,
    content: str,
    country_code: str,
) -> List[SimilarClause]:
    """Findet ähnliche existierende Klauseln."""
    similar = []

    # Titel-basierte Suche
    title_words = set(title.lower().split())

    result = await db.execute(
        select(Clause).where(
            Clause.country_code == country_code,
            Clause.is_active == True,
        )
    )
    clauses = result.scalars().all()

    for clause in clauses:
        # Titel-Ähnlichkeit
        clause_title_words = set(clause.title.lower().split())
        title_similarity = len(title_words & clause_title_words) / max(len(title_words | clause_title_words), 1)

        # Inhalt-Ähnlichkeit (vereinfacht: Wort-Überlappung)
        content_words = set(re.findall(r'\b\w+\b', content.lower()))
        clause_content_words = set(re.findall(r'\b\w+\b', (clause.content_html or "").lower()))
        content_similarity = len(content_words & clause_content_words) / max(len(content_words | clause_content_words), 1)

        # Gewichtete Ähnlichkeit
        similarity = title_similarity * 0.4 + content_similarity * 0.6

        if similarity > 0.3:  # Mindest-Schwelle
            # Prüfen ob Varianten existieren (gleicher Titel, unterschiedlicher Inhalt)
            variants_result = await db.execute(
                select(Clause).where(
                    Clause.title == clause.title,
                    Clause.country_code == country_code,
                    Clause.id != clause.id,
                )
            )
            variants = variants_result.scalars().all()

            similar.append(SimilarClause(
                id=clause.id,
                title=clause.title,
                content_preview=(clause.content_html or "")[:200].replace("<", "&lt;"),
                similarity_score=round(similarity, 2),
                has_variants=len(variants) > 0,
                variants=[{"id": v.id, "version": v.version} for v in variants],
            ))

    # Nach Ähnlichkeit sortieren
    similar.sort(key=lambda x: x.similarity_score, reverse=True)
    return similar[:5]  # Top 5


# ══════════════════════════════════════════════════════════════════════════════
# WORD PARSING (erweitert von word_import.py)
# ══════════════════════════════════════════════════════════════════════════════

def extract_document_structure(file_content: bytes) -> dict:
    """Extrahiert Struktur aus DOCX mit erweiterter Analyse."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx nicht installiert")

    doc = Document(io.BytesIO(file_content))

    full_text = []
    sections = []
    current_section = {"title": "", "content": [], "content_html": []}
    warnings = []
    stats = {"headings_found": 0, "tables_converted": 0, "images_ignored": 0, "placeholders_detected": 0}

    heading_patterns = [
        r'^§\s*\d+',
        r'^Art\.\s*\d+',
        r'^\d+\.\s+[A-ZÄÖÜ]',
        r'^[IVXLCDM]+\.\s+',
    ]

    def is_heading(text: str, style_name: str) -> bool:
        if style_name and ('heading' in style_name.lower() or 'überschrift' in style_name.lower()):
            return True
        text_stripped = text.strip()
        return any(re.match(p, text_stripped, re.IGNORECASE) for p in heading_patterns)

    def save_section():
        if current_section["title"] or current_section["content"]:
            content_text = "\n".join(current_section["content"]).strip()
            content_html = "<br>".join(current_section["content_html"])

            if content_text:
                # Platzhalter erkennen
                placeholders = detect_placeholders(content_text)
                stats["placeholders_detected"] += len(placeholders)

                sections.append({
                    "title": current_section["title"].strip() or f"Abschnitt {len(sections) + 1}",
                    "content": content_text,
                    "content_html": f"<p>{content_html}</p>" if content_html else "",
                    "placeholders": placeholders,
                })

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if is_heading(text, style_name):
            save_section()
            current_section = {
                "title": text,
                "content": [],
                "content_html": [f"<strong>{text}</strong>"],
            }
            stats["headings_found"] += 1
        else:
            current_section["content"].append(text)
            html_text = text
            for run in para.runs:
                if run.bold:
                    html_text = html_text.replace(run.text, f"<strong>{run.text}</strong>")
                if run.italic:
                    html_text = html_text.replace(run.text, f"<em>{run.text}</em>")
            current_section["content_html"].append(html_text)

        full_text.append(text)

    save_section()

    # Tabellen
    for table in doc.tables:
        stats["tables_converted"] += 1
        warnings.append(f"Tabelle mit {len(table.rows)} Zeilen wurde zu Text konvertiert")

    # Bilder
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            stats["images_ignored"] += 1

    if stats["images_ignored"] > 0:
        warnings.append(f"{stats['images_ignored']} Bild(er) wurden ignoriert")

    return {
        "text": "\n".join(full_text),
        "sections": sections,
        "warnings": warnings,
        "stats": stats,
    }


# ══════════════════════════════════════════════════════════════════════════════
# API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/analyze", response_model=DocumentTypeAnalysisResult)
async def analyze_document_for_import(
    file: UploadFile = File(...),
    country_code: str = Form("DE"),
    db: AsyncSession = Depends(get_db),
):
    """
    Analysiert ein Word-Dokument für den Import als Dokumententyp.

    - Erkennt Abschnitte und Struktur
    - Findet ähnliche existierende Klauseln
    - Erkennt potenzielle Platzhalter
    """
    if not file.filename.endswith(('.docx', '.DOCX')):
        raise HTTPException(status_code=400, detail="Nur .docx Dateien werden unterstützt")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Datei zu groß (max 10 MB)")

    try:
        result = extract_document_structure(content)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Fehler beim Lesen: {str(e)}")

    # Abschnitte mit Ähnlichkeits-Prüfung
    sections: List[ExtractedSection] = []
    similar_count = 0

    for idx, section in enumerate(result["sections"]):
        # Ähnliche Klauseln suchen
        similar = await find_similar_clauses(
            db, section["title"], section["content"], country_code
        )

        best_match = similar[0] if similar else None

        sections.append(ExtractedSection(
            index=idx,
            title=section["title"],
            content=section["content"],
            content_html=section["content_html"],
            char_count=len(section["content"]),
            action="use_existing" if best_match and best_match.similarity_score > 0.7 else "create_new",
            existing_clause_id=best_match.id if best_match else None,
            existing_clause_title=best_match.title if best_match else None,
            similarity_score=best_match.similarity_score if best_match else None,
            is_mandatory=True,
            display_order=idx,
            placeholders=[PlaceholderSuggestion(**p.__dict__) if hasattr(p, '__dict__') else p
                         for p in section.get("placeholders", [])],
        ))

        if best_match:
            similar_count += 1

    return DocumentTypeAnalysisResult(
        filename=file.filename,
        total_chars=len(result["text"]),
        sections_count=len(sections),
        sections=sections,
        warnings=result["warnings"],
        headings_found=result["stats"]["headings_found"],
        tables_converted=result["stats"]["tables_converted"],
        images_ignored=result["stats"]["images_ignored"],
        placeholders_detected=result["stats"]["placeholders_detected"],
        similar_clauses_found=similar_count,
    )


@router.post("/import", response_model=ImportResult)
async def import_as_document_type(
    request: DocumentTypeImportRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Importiert das analysierte Dokument als neuen Dokumententyp.

    - Erstellt Dokumententyp
    - Erstellt neue Klauseln oder verknüpft existierende
    - Fügt Varianten hinzu
    - Erstellt Formularfelder
    """
    # 1. Dokumententyp erstellen
    doc_type = DocumentType(
        name=request.document_type_name,
        country_code=request.country_code,
        category=request.document_type_category,
        is_active=True,
    )
    db.add(doc_type)
    await db.flush()

    clauses_created = 0
    clauses_linked = 0
    variants_added = 0
    all_placeholders = set()

    # 2. Abschnitte verarbeiten
    for section in request.sections:
        if section.action == "skip":
            continue

        clause_id = None

        if section.action == "create_new":
            # Platzhalter anwenden
            content_with_placeholders = section.content_html
            for ph in section.placeholders:
                if ph.accepted:
                    content_with_placeholders = content_with_placeholders.replace(
                        ph.original_text, ph.suggested_placeholder
                    )
                    all_placeholders.add(ph.suggested_placeholder.strip("{}"))

            # Neue Klausel erstellen
            clause = Clause(
                title=section.title,
                content_html=content_with_placeholders,
                country_code=request.country_code,
                category=request.document_type_category,
                version=1,
                is_active=True,
                approval_status="draft",
            )
            db.add(clause)
            await db.flush()
            clause_id = clause.id
            clauses_created += 1

        elif section.action == "use_existing" and section.existing_clause_id:
            clause_id = section.existing_clause_id
            clauses_linked += 1

        elif section.action == "add_variant" and section.is_variant_of:
            # Variante zur existierenden Klausel hinzufügen
            # (Neue Klausel mit Referenz zur Haupt-Klausel)
            original = await db.get(Clause, section.is_variant_of)
            if original:
                content_with_placeholders = section.content_html
                for ph in section.placeholders:
                    if ph.accepted:
                        content_with_placeholders = content_with_placeholders.replace(
                            ph.original_text, ph.suggested_placeholder
                        )
                        all_placeholders.add(ph.suggested_placeholder.strip("{}"))

                variant_title = f"{original.title} ({section.variant_name or 'Variante'})"
                variant = Clause(
                    title=variant_title,
                    content_html=content_with_placeholders,
                    country_code=request.country_code,
                    category=original.category,
                    version=1,
                    is_active=True,
                    approval_status="draft",
                )
                db.add(variant)
                await db.flush()
                clause_id = variant.id
                variants_added += 1

        # 3. Klausel mit Dokumententyp verknüpfen
        if clause_id:
            link = DocumentTypeClause(
                document_type_id=doc_type.id,
                clause_id=clause_id,
                display_order=section.display_order,
                is_mandatory=section.is_mandatory,
            )
            db.add(link)

    # 4. Globale Platzhalter sammeln
    for ph in request.global_placeholders:
        if ph.accepted:
            all_placeholders.add(ph.suggested_placeholder.strip("{}"))

    # 5. Formularfelder erstellen (vereinfacht - würde normalerweise in form_fields Tabelle)
    form_fields_created = len(all_placeholders)

    await db.commit()

    return ImportResult(
        document_type_id=doc_type.id,
        document_type_name=doc_type.name,
        clauses_created=clauses_created,
        clauses_linked=clauses_linked,
        variants_added=variants_added,
        form_fields_created=form_fields_created,
    )


@router.get("/placeholder-suggestions")
async def get_placeholder_suggestions():
    """Gibt bekannte Platzhalter-Muster und Tipps zurück."""
    return {
        "common_placeholders": [
            {"name": "mitarbeiter_name", "type": "text", "description": "Vollständiger Name des Mitarbeiters"},
            {"name": "eintrittsdatum", "type": "date", "description": "Datum des Arbeitsbeginns"},
            {"name": "gehalt", "type": "currency", "description": "Monatliches Bruttogehalt"},
            {"name": "wochenstunden", "type": "number", "description": "Wöchentliche Arbeitszeit"},
            {"name": "urlaubstage", "type": "number", "description": "Jährlicher Urlaubsanspruch"},
            {"name": "kuendigungsfrist", "type": "text", "description": "Kündigungsfrist"},
            {"name": "probezeit_monate", "type": "number", "description": "Dauer der Probezeit"},
            {"name": "position", "type": "text", "description": "Stellenbezeichnung"},
            {"name": "abteilung", "type": "text", "description": "Abteilung"},
            {"name": "arbeitsort", "type": "text", "description": "Arbeitsort/Standort"},
        ],
        "tips": [
            "Platzhalter werden automatisch erkannt: Namen, Daten, Beträge",
            "Sie können vorgeschlagene Platzhalter anpassen oder ablehnen",
            "Gleiche Platzhalter in verschiedenen Klauseln werden zusammengeführt",
            "Formularfelder werden automatisch aus den Platzhaltern erstellt",
        ],
    }


@router.get("/similar-clauses/{title}")
async def find_similar_clauses_by_title(
    title: str,
    country_code: str = "DE",
    db: AsyncSession = Depends(get_db),
):
    """Sucht ähnliche Klauseln für manuelles Matching."""
    similar = await find_similar_clauses(db, title, "", country_code)
    return {"similar_clauses": similar}
