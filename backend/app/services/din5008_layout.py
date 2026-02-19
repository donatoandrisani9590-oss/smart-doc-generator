"""
DIN 5008 Layout Builder -- Erstellt professionelles Dokumentlayout nach DIN 5008.

Verwendet bei fehlendem Briefpapier (kein UserTemplate vorhanden).
Generiert Kopfzeile mit Firmendaten, Logo, Fusszeile, korrekte Raender.

DIN 5008 Typ B (Standard):
- Oberer Rand: 2,7 cm (Kopfzeile)
- Linker Rand: 2,5 cm
- Rechter Rand: 2,0 cm (mind. 1,0 cm)
- Unterer Rand: 2,0 cm
- Schriftgroesse: 11pt (Arial/Helvetica)
- Zeilenabstand: Einfach oder 1,15-fach
"""
import logging
import os
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class DIN5008LayoutBuilder:
    """Creates a DIN 5008-compliant DOCX document layout as fallback when no Briefpapier is uploaded."""

    def __init__(
        self,
        company_name: str = "",
        company_address: str = "",
        company_contact: str = "",
        logo_path: Optional[str] = None,
        primary_color: str = "#243186",
        font_family: str = "Arial",
        font_size_pt: int = 11,
        footer_lines: Optional[list] = None,
    ):
        """
        Initialize the DIN 5008 layout builder.

        Args:
            company_name: Firmenname fuer die Kopfzeile
            company_address: Adresszeile (Strasse, PLZ Ort)
            company_contact: Kontaktinformationen (E-Mail, Telefon)
            logo_path: Pfad zum Firmenlogo (optional)
            primary_color: Primaerfarbe als Hex (z.B. "#243186")
            font_family: Schriftart (Standard: Arial)
            font_size_pt: Schriftgroesse in Punkt (Standard: 11)
            footer_lines: Optionale Fusszeilen-Texte (z.B. Handelsregister, USt-IdNr.)
        """
        self.company_name = company_name
        self.company_address = company_address
        self.company_contact = company_contact
        self.logo_path = logo_path
        self.primary_color = primary_color
        self.font_family = font_family
        self.font_size_pt = font_size_pt
        self.footer_lines = footer_lines or []

    def apply_layout(self, doc: Document) -> None:
        """Apply DIN 5008 Typ B layout to the document."""
        self._set_margins(doc)
        self._set_default_font(doc)
        self._configure_heading_styles(doc)
        self._add_header(doc)
        self._add_footer(doc)

    def _set_margins(self, doc: Document) -> None:
        """Set DIN 5008 Typ B margins on all sections."""
        for section in doc.sections:
            section.top_margin = Cm(2.7)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.0)

    def _set_default_font(self, doc: Document) -> None:
        """Set default font and paragraph format for the Normal style."""
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_family
        font.size = Pt(self.font_size_pt)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)  # Near-black for readability

        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

        # Ensure the font is also set for East Asian / Complex Script fallback
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), self.font_family)
        rFonts.set(qn('w:hAnsi'), self.font_family)

    def _configure_heading_styles(self, doc: Document) -> None:
        """Configure Heading 1 and Heading 2 styles for consistent typography."""
        r, g, b = self._parse_hex_color(self.primary_color)

        # Heading 1: 14pt, bold, primary color
        try:
            h1_style = doc.styles['Heading 1']
            h1_style.font.name = self.font_family
            h1_style.font.size = Pt(14)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(r, g, b)
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            logger.debug("Heading 1 style not available in template")

        # Heading 2: 12pt, bold, primary color
        try:
            h2_style = doc.styles['Heading 2']
            h2_style.font.name = self.font_family
            h2_style.font.size = Pt(12)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(r, g, b)
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(4)
        except KeyError:
            logger.debug("Heading 2 style not available in template")

    def _parse_hex_color(self, hex_color: str) -> tuple:
        """Parse a hex color string to (r, g, b) tuple. Falls back to primary blue."""
        try:
            hex_color = hex_color.lstrip('#')
            if len(hex_color) == 6:
                return (
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16),
                )
        except (ValueError, IndexError):
            pass
        # Fallback: Corporate blue
        return (0x24, 0x31, 0x86)

    def _add_header(self, doc: Document) -> None:
        """Add company header to the document (DIN 5008 compliant)."""
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing header content
        for p in header.paragraphs:
            p.clear()

        r, g, b = self._parse_hex_color(self.primary_color)

        # If logo exists, add it right-aligned above company name
        logo_added = self._try_add_logo(header)

        # Company name - right-aligned, small, bold, in primary color
        if self.company_name:
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            # If logo was added, use the existing first paragraph or add a new one
            if logo_added:
                p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(self.company_name)
            run.font.size = Pt(9)
            run.font.name = self.font_family
            run.font.color.rgb = RGBColor(r, g, b)
            run.bold = True

        # Company address + contact - right-aligned, smaller, gray
        if self.company_address or self.company_contact:
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            details = []
            if self.company_address:
                details.append(self.company_address)
            if self.company_contact:
                details.append(self.company_contact)
            run = p.add_run(" \u00b7 ".join(details))
            run.font.size = Pt(8)
            run.font.name = self.font_family
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Add thin separator line below header
        self._add_border_paragraph(header, border_side='bottom')

    def _try_add_logo(self, header) -> bool:
        """Try to add a company logo to the header. Returns True if successful."""
        if not self.logo_path:
            return False

        # Check multiple possible paths for the logo
        possible_paths = [self.logo_path]
        if not self.logo_path.startswith('/'):
            possible_paths.append(f"/app/storage/logos/{self.logo_path}")

        for lp in possible_paths:
            if os.path.exists(lp):
                try:
                    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(lp, width=Cm(3.5))  # Compact logo in header
                    return True
                except Exception as e:
                    logger.warning(f"Konnte Logo nicht hinzufuegen ({lp}): {e}")
        return False

    def _add_footer(self, doc: Document) -> None:
        """Add professional footer with page numbers and optional company info."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer
        for p in footer.paragraphs:
            p.clear()

        # Add separator line above footer
        self._add_border_paragraph(footer, border_side='top', is_first=True)

        # Optional footer lines (Handelsregister, USt-IdNr., etc.)
        if self.footer_lines:
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text = " \u00b7 ".join([line for line in self.footer_lines if line])
            run = p.add_run(footer_text)
            run.font.size = Pt(7)
            run.font.name = self.font_family
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Page number: "Seite X von Y"
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "Seite " text
        run = p.add_run("Seite ")
        run.font.size = Pt(8)
        run.font.name = self.font_family
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Current page number (PAGE field)
        self._add_field(run._element, 'PAGE')

        # " von " text
        run2 = p.add_run(" von ")
        run2.font.size = Pt(8)
        run2.font.name = self.font_family
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Total pages (NUMPAGES field)
        self._add_field(run2._element, 'NUMPAGES')

    def _add_border_paragraph(
        self,
        container,
        border_side: str = 'bottom',
        is_first: bool = False,
    ) -> None:
        """Add a paragraph with a thin border line (used for header/footer separators)."""
        if is_first and container.paragraphs:
            p = container.paragraphs[0]
        else:
            p = container.add_paragraph()

        if border_side == 'bottom':
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
        else:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)

        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        border_el = OxmlElement(f'w:{border_side}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')  # 0.5pt line
        border_el.set(qn('w:space'), '1')
        border_el.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border_el)
        pPr.append(pBdr)

    @staticmethod
    def _add_field(run_element, field_code: str) -> None:
        """Add a Word field (e.g., PAGE, NUMPAGES) to a run element."""
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run_element.append(fldChar_begin)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code
        run_element.append(instrText)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run_element.append(fldChar_end)
