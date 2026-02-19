"""
HTML to DOCX Converter with full table and nested list support.

This service converts HTML content to python-docx elements, preserving:
- Tables with headers, cell formatting, and borders
- Nested lists (ul/ol) with proper indentation levels
- Inline formatting (bold, italic, underline)
- Paragraphs with alignment
- Headings with proper styling

Replaces the simpler html_to_docx_paragraph functions in generation.py
"""
from __future__ import annotations
from typing import Any, Optional, Iterator
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import logging

logger = logging.getLogger(__name__)

# Maximum nesting depth for lists (matches frontend ConditionBuilder limit)
MAX_LIST_DEPTH = 3


class HtmlToDocxConverter:
    """
    Converts HTML content to DOCX elements with full formatting support.

    Usage:
        converter = HtmlToDocxConverter()
        converter.convert(doc, html_content, form_data, country_code)
    """

    def __init__(self, font_size_pt: int = 11, font_family: str = "Arial"):
        """
        Initialize converter with default styling.

        Args:
            font_size_pt: Default font size in points
            font_family: Default font family name
        """
        self.font_size_pt = font_size_pt
        self.font_family = font_family

    def convert(
        self,
        doc: Document,
        html_content: str,
        form_data: dict,
        country_code: str = "DE"
    ) -> None:
        """
        Convert HTML content and add it to the document.

        Args:
            doc: The python-docx Document object
            html_content: HTML string to convert
            form_data: Dictionary of placeholder values
            country_code: Country code for localization (DE/IT)
        """
        if not html_content:
            return

        # Replace placeholders first
        rendered_html = self._render_placeholders(html_content, form_data, country_code)

        # Parse HTML
        soup = BeautifulSoup(rendered_html, 'html.parser')

        # Process all top-level elements
        self._process_elements(doc, soup.children, depth=0)

    def _render_placeholders(
        self,
        content: str,
        form_data: dict,
        country_code: str
    ) -> str:
        """Replace {{ placeholder }} and [placeholder] with values."""
        from app.services.preview import render_placeholders
        return render_placeholders(content, form_data, country_code)

    def _process_elements(
        self,
        doc: Document,
        elements: Iterator,
        depth: int = 0,
        parent_cell: Any = None
    ) -> None:
        """
        Recursively process HTML elements.

        Args:
            doc: Document or cell to add content to
            elements: Iterator of HTML elements
            depth: Current nesting depth (for lists)
            parent_cell: If processing inside a table cell
        """
        container = parent_cell if parent_cell else doc

        for element in elements:
            if isinstance(element, NavigableString):
                # Plain text node - skip if whitespace only
                text = str(element).strip()
                if text:
                    para = container.add_paragraph()
                    self._add_run(para, text)
                continue

            if not isinstance(element, Tag):
                continue

            tag_name = element.name.lower() if element.name else ""

            # Table
            if tag_name == 'table':
                self._add_table(container, element)

            # Lists
            elif tag_name in ('ul', 'ol'):
                self._add_list(container, element, depth=0, is_ordered=(tag_name == 'ol'))

            # Headings
            elif tag_name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                self._add_heading(container, element, tag_name)

            # Paragraphs and divs (check for page-break class first)
            elif tag_name in ('p', 'div'):
                css_class = element.get('class', [])
                if isinstance(css_class, str):
                    css_class = css_class.split()
                if 'page-break' in css_class or 'pagebreak' in css_class:
                    self._add_page_break(container)
                else:
                    self._add_paragraph(container, element)

            # Blockquote
            elif tag_name == 'blockquote':
                self._add_blockquote(container, element)

            # Horizontal rule (check for page-break class)
            elif tag_name == 'hr':
                css_class = element.get('class', [])
                if isinstance(css_class, str):
                    css_class = css_class.split()
                if 'page-break' in css_class or 'pagebreak' in css_class:
                    self._add_page_break(container)
                else:
                    self._add_horizontal_rule(container)

            # Line break at top level - add empty paragraph
            elif tag_name == 'br':
                container.add_paragraph()

            # Pre/code blocks
            elif tag_name in ('pre', 'code'):
                self._add_code_block(container, element)

            # Span with potential styling
            elif tag_name == 'span':
                # Process inline content
                para = container.add_paragraph()
                self._add_inline_content(para, element)

            # Unknown elements - try to process children
            else:
                if element.children:
                    self._process_elements(doc, element.children, depth, parent_cell)

    def _add_table(self, container: Any, table_el: Tag) -> None:
        """
        Convert HTML table to DOCX table.

        Handles:
        - thead/tbody structure
        - th (header) cells with bold styling
        - colspan/rowspan (basic support)
        - Cell content with inline formatting
        """
        rows = table_el.find_all('tr')
        if not rows:
            return

        # Determine column count (max cells in any row)
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            col_count = 0
            for cell in cells:
                try:
                    col_count += int(cell.get('colspan', 1))
                except (ValueError, TypeError):
                    col_count += 1
            max_cols = max(max_cols, col_count)

        if max_cols == 0:
            return

        # Create table
        table = container.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        # Process rows
        for row_idx, row_el in enumerate(rows):
            cells = row_el.find_all(['td', 'th'])
            col_idx = 0

            for cell_el in cells:
                if col_idx >= max_cols:
                    break

                cell = table.cell(row_idx, col_idx)
                is_header = cell_el.name == 'th'

                # Clear default paragraph
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                # Process cell content
                self._process_cell_content(cell, cell_el, is_header)

                # Handle colspan
                try:
                    colspan = int(cell_el.get('colspan', 1))
                except (ValueError, TypeError):
                    colspan = 1
                if colspan > 1 and col_idx + colspan <= max_cols:
                    for merge_idx in range(1, colspan):
                        cell.merge(table.cell(row_idx, col_idx + merge_idx))

                col_idx += colspan

        # Add table borders
        self._set_table_borders(table)

    def _process_cell_content(
        self,
        cell: Any,
        cell_el: Tag,
        is_header: bool = False
    ) -> None:
        """Process content inside a table cell."""
        # Get or create first paragraph
        if cell.paragraphs:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()

        # Check for nested block elements
        block_elements = cell_el.find_all(['p', 'ul', 'ol', 'table'], recursive=False)

        if block_elements:
            # Has block-level content - process recursively
            for idx, block in enumerate(block_elements):
                if idx > 0:
                    para = cell.add_paragraph()

                if block.name == 'table':
                    # Nested table - convert to text (Word nested tables are complex)
                    self._add_inline_content(para, block)
                elif block.name in ('ul', 'ol'):
                    self._add_list_to_cell(cell, block, is_ordered=(block.name == 'ol'))
                else:
                    self._add_inline_content(para, block)
        else:
            # Simple content
            self._add_inline_content(para, cell_el)

        # Apply header styling
        if is_header:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    def _add_list(
        self,
        container: Any,
        list_el: Tag,
        depth: int = 0,
        is_ordered: bool = False
    ) -> None:
        """
        Convert HTML list to DOCX list paragraphs with proper nesting.

        Args:
            container: Document or cell to add to
            list_el: The <ul> or <ol> element
            depth: Current nesting level (0-based)
            is_ordered: True for numbered list, False for bullet
        """
        if depth > MAX_LIST_DEPTH:
            depth = MAX_LIST_DEPTH

        for li in list_el.find_all('li', recursive=False):
            # Create list paragraph
            para = container.add_paragraph()

            # Set list style and indentation
            self._apply_list_style(para, depth, is_ordered)

            # Add inline content from <li> (excluding nested lists)
            self._add_list_item_content(para, li)

            # Process nested lists
            for nested in li.find_all(['ul', 'ol'], recursive=False):
                self._add_list(
                    container,
                    nested,
                    depth=depth + 1,
                    is_ordered=(nested.name == 'ol')
                )

    def _add_list_to_cell(
        self,
        cell: Any,
        list_el: Tag,
        is_ordered: bool = False
    ) -> None:
        """Add a list inside a table cell."""
        for li in list_el.find_all('li', recursive=False):
            para = cell.add_paragraph()

            # Simple bullet/number prefix for cell lists
            prefix = "• " if not is_ordered else f"{li.parent.index(li) + 1}. "
            run = para.add_run(prefix)

            # Add content
            self._add_list_item_content(para, li)

    def _add_list_item_content(self, para: Any, li: Tag) -> None:
        """Add content from a <li> element, excluding nested lists."""
        for child in li.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    self._add_run(para, text)
            elif isinstance(child, Tag):
                # Skip nested lists - they're processed separately
                if child.name in ('ul', 'ol'):
                    continue
                # Process inline elements
                self._add_inline_element(para, child)

    def _apply_list_style(
        self,
        para: Any,
        depth: int,
        is_ordered: bool
    ) -> None:
        """Apply list styling with proper indentation."""
        # Indentation: 0.5 inch per level
        left_indent = Cm(1.27 * (depth + 1))  # ~0.5 inch per level
        hanging = Cm(0.635)  # ~0.25 inch hanging indent

        para.paragraph_format.left_indent = left_indent
        para.paragraph_format.first_line_indent = -hanging

        # Try to use built-in list style, fall back to manual bullet
        try:
            style_name = 'List Number' if is_ordered else 'List Bullet'
            para.style = style_name
        except KeyError:
            # Manual bullet/number if style not available
            if is_ordered:
                # For ordered lists, we'd need to track numbering
                prefix = "1. "  # Simplified - real implementation would track
            else:
                bullets = ["•", "◦", "▪"]
                prefix = bullets[min(depth, len(bullets) - 1)] + " "

            # Prepend to first run or add new
            if para.runs:
                para.runs[0].text = prefix + para.runs[0].text
            else:
                para.add_run(prefix)

    def _add_heading(self, container: Any, element: Tag, tag_name: str) -> None:
        """
        Add a heading using proper Word heading styles (Heading 1, Heading 2, etc.).

        Using built-in heading styles ensures:
        - Automatic Table of Contents (TOC) support
        - Consistent typography across the document
        - Proper outline level for navigation
        """
        level = int(tag_name[1])  # h1 -> 1, h2 -> 2, etc.

        # Use proper Word heading styles for levels 1-3, fall back to manual for 4-6
        if level <= 3:
            try:
                heading_style = f'Heading {level}'
                para = container.add_paragraph(style=heading_style)
                self._add_inline_content(para, element)
                # Ensure correct font family on heading runs
                for run in para.runs:
                    run.font.name = self.font_family
                return
            except KeyError:
                # Heading style not available -- fall through to manual formatting
                logger.debug(f"Word style 'Heading {level}' nicht verfuegbar, verwende manuelle Formatierung")

        # Manual heading formatting (levels 4-6 or when styles are unavailable)
        para = container.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        sizes = {1: 16, 2: 14, 3: 12, 4: 12, 5: 11, 6: 11}
        font_size = sizes.get(level, 11)

        self._add_inline_content(para, element)

        # Apply heading formatting to all runs
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = self.font_family

        # Add space before heading
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)

    def _add_paragraph(self, container: Any, element: Tag) -> None:
        """Add a regular paragraph with alignment and page-break support."""
        style_attr = element.get('style', '')

        # Handle CSS page-break-before: always / page-break-after: always
        if 'page-break-before' in style_attr and 'always' in style_attr:
            self._add_page_break(container)

        para = container.add_paragraph()

        # Check for alignment style
        if 'text-align: center' in style_attr or 'text-align:center' in style_attr:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in style_attr or 'text-align:right' in style_attr:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'text-align: justify' in style_attr or 'text-align:justify' in style_attr:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Default to justified

        self._add_inline_content(para, element)

        # Handle page-break-after
        if 'page-break-after' in style_attr and 'always' in style_attr:
            self._add_page_break(container)

    def _add_blockquote(self, container: Any, element: Tag) -> None:
        """Add a blockquote with indentation."""
        para = container.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.right_indent = Cm(1.0)

        self._add_inline_content(para, element)

        # Italic style for quotes
        for run in para.runs:
            run.italic = True

    def _add_code_block(self, container: Any, element: Tag) -> None:
        """Add a preformatted code block."""
        para = container.add_paragraph()

        text = element.get_text()
        run = para.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    def _add_horizontal_rule(self, container: Any) -> None:
        """Add a horizontal rule (as a thin paragraph with border)."""
        para = container.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # Add bottom border to simulate HR
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_page_break(self, container: Any) -> None:
        """
        Add an explicit page break to the document.

        Triggered by:
        - <div class="page-break">
        - <hr class="page-break">
        - <p class="pagebreak">
        """
        para = container.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _add_inline_content(self, para: Any, element: Tag) -> None:
        """
        Process inline content within a paragraph.

        Handles: strong/b, em/i, u, span, a, br, sub, sup
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:  # Keep whitespace for inline flow
                    self._add_run(para, text)
            elif isinstance(child, Tag):
                self._add_inline_element(para, child)

    def _add_inline_element(self, para: Any, element: Tag) -> None:
        """Process a single inline element."""
        tag = element.name.lower() if element.name else ""

        if tag == 'br':
            # Line break within paragraph
            if para.runs:
                para.runs[-1].add_break()
            else:
                para.add_run().add_break()

        elif tag in ('strong', 'b'):
            run = para.add_run(element.get_text())
            run.bold = True

        elif tag in ('em', 'i'):
            run = para.add_run(element.get_text())
            run.italic = True

        elif tag == 'u':
            run = para.add_run(element.get_text())
            run.underline = True

        elif tag == 'sub':
            run = para.add_run(element.get_text())
            run.font.subscript = True

        elif tag == 'sup':
            run = para.add_run(element.get_text())
            run.font.superscript = True

        elif tag == 'a':
            # Hyperlink - just show text (real hyperlinks need more complex handling)
            run = para.add_run(element.get_text())
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 255)

        elif tag == 'span':
            # Check for style
            style = element.get('style', '')
            run = para.add_run(element.get_text())

            if 'font-weight: bold' in style or 'font-weight:bold' in style:
                run.bold = True
            if 'font-style: italic' in style or 'font-style:italic' in style:
                run.italic = True
            if 'text-decoration: underline' in style or 'text-decoration:underline' in style:
                run.underline = True

            # Color extraction
            color_match = re.search(r'color:\s*#([0-9a-fA-F]{6})', style)
            if color_match:
                hex_color = color_match.group(1)
                run.font.color.rgb = RGBColor(
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16)
                )

        else:
            # Unknown inline element - just get text
            text = element.get_text()
            if text:
                self._add_run(para, text)

    def _add_run(self, para: Any, text: str) -> None:
        """Add a text run with default formatting."""
        run = para.add_run(text)
        run.font.size = Pt(self.font_size_pt)
        run.font.name = self.font_family

    def _set_table_borders(self, table: Any) -> None:
        """Apply borders to all cells in the table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)
        tbl.tblPr = tblPr
