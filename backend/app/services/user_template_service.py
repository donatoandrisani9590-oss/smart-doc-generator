"""
User Template Service - Inject content into user-uploaded DOCX templates.

Takes a user's branded DOCX template (with header/footer/logo) and injects
generated content into the body while preserving all layout elements.
"""
import logging
from pathlib import Path
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.html_converter import HtmlToDocxConverter

logger = logging.getLogger(__name__)

# Storage paths
BASE_DIR = Path(__file__).parent.parent.parent
USER_TEMPLATE_STORAGE = BASE_DIR / "storage" / "user-templates"
OUTPUT_DIR = BASE_DIR / "storage" / "generated"


def inject_content_into_template(
    template_path: Path,
    clauses: list[dict],
    form_data: dict,
    country_code: str = "DE",
    design_settings: Optional[dict] = None,
    output_path: Optional[Path] = None,
) -> Path:
    """
    Open a user template, clear its body, and inject clause content.

    The template's headers, footers, images, styles, and page setup are
    fully preserved. Only the body paragraphs are replaced.

    Args:
        template_path: Path to the user's DOCX template
        clauses: List of clause dicts with 'content' (HTML) and 'title'
        form_data: Form data for placeholder replacement
        country_code: Country code for localization (DE/IT)
        design_settings: Optional design settings (font_size, font_family)
        output_path: Where to save the result

    Returns:
        Path to the generated DOCX
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template nicht gefunden: {template_path}")

    # Open the template
    doc = Document(str(template_path))

    # Detect template's font from Normal style or first paragraph
    template_font = None
    template_font_size = None
    try:
        style = doc.styles['Normal']
        if style.font and style.font.name:
            template_font = style.font.name
        if style.font and style.font.size:
            template_font_size = style.font.size.pt
    except Exception:
        pass

    # Fallback: check first paragraph
    if not template_font:
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.name:
                    template_font = run.font.name
                    break
            if template_font:
                break

    # Use template font or design settings or default
    font_family = template_font or (design_settings or {}).get('font_family', 'Arial')
    font_size = template_font_size or (design_settings or {}).get('font_size_pt', 12)

    # Clear the body: remove all existing paragraphs
    # We must keep at least one paragraph (python-docx requirement)
    body = doc.element.body

    # Remove all paragraphs and tables from body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            body.remove(child)

    # Now add content using the HTML converter
    html_converter = HtmlToDocxConverter(
        font_size_pt=font_size,
        font_family=font_family,
    )

    for i, clause in enumerate(clauses):
        content = clause.get('content', '')
        if not content:
            continue

        # Add clause content (HTML -> DOCX paragraphs)
        html_converter.convert(doc, content, form_data, country_code)

        # Add spacing between clauses (except after the last one)
        if i < len(clauses) - 1:
            doc.add_paragraph()

    # If no clauses were added, add an empty paragraph (python-docx needs at least one)
    if not clauses:
        doc.add_paragraph()

    # Save the document
    if output_path is None:
        output_path = OUTPUT_DIR / "generated_from_template.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"Document generated from user template: {output_path}")

    return output_path


def inject_freetext_into_template(
    template_path: Path,
    html_content: str,
    form_data: dict,
    country_code: str = "DE",
    design_settings: Optional[dict] = None,
    output_path: Optional[Path] = None,
) -> Path:
    """
    Inject free-form HTML content (e.g. AI-generated) into a user template.

    Similar to inject_content_into_template but takes a single HTML string
    instead of a list of clauses.

    Args:
        template_path: Path to the user's DOCX template
        html_content: HTML content to inject (e.g. AI-generated letter)
        form_data: Form data for placeholder replacement
        country_code: Country code for localization
        design_settings: Optional design settings
        output_path: Where to save the result

    Returns:
        Path to the generated DOCX
    """
    clauses = [{"content": html_content, "title": ""}]
    return inject_content_into_template(
        template_path=template_path,
        clauses=clauses,
        form_data=form_data,
        country_code=country_code,
        design_settings=design_settings,
        output_path=output_path,
    )


def get_user_template_path(user_id: int, stored_filename: str) -> Path:
    """Get the full file path for a user template."""
    return USER_TEMPLATE_STORAGE / str(user_id) / stored_filename


def extract_template_zones(template_path: Path) -> dict:
    """
    Extrahiert Header-/Footer-Bilder und Text aus einer DOCX-Vorlage.

    Gibt base64-kodierte Bilder und Textinhalte zurück, die für
    die Canvas-Vorschau im Frontend verwendet werden können.

    Args:
        template_path: Pfad zur DOCX-Vorlage

    Returns:
        Dict mit header_images, footer_images, header_text, footer_text, page_margins
    """
    import base64
    from docx.shared import Emu

    doc = Document(str(template_path))
    result = {
        "header_images": [],
        "footer_images": [],
        "header_text": "",
        "footer_text": "",
        "page_margins": {
            "top": 0,
            "bottom": 0,
            "left": 0,
            "right": 0,
        },
    }

    # Seitenraender extrahieren (EMU -> cm)
    for section in doc.sections:
        result["page_margins"] = {
            "top": round(section.top_margin / 914400 * 2.54, 2) if section.top_margin else 1.27,
            "bottom": round(section.bottom_margin / 914400 * 2.54, 2) if section.bottom_margin else 1.27,
            "left": round(section.left_margin / 914400 * 2.54, 2) if section.left_margin else 1.27,
            "right": round(section.right_margin / 914400 * 2.54, 2) if section.right_margin else 1.27,
        }

        # Header-Inhalt extrahieren
        header = section.header
        if header and header.paragraphs:
            texts = []
            for para in header.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            result["header_text"] = "\n".join(texts)

        # Footer-Inhalt extrahieren
        footer = section.footer
        if footer and footer.paragraphs:
            texts = []
            for para in footer.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            result["footer_text"] = "\n".join(texts)

    # Bilder aus dem DOCX-Paket extrahieren
    try:
        import zipfile
        with zipfile.ZipFile(str(template_path), 'r') as zf:
            # Alle Bilddateien finden
            image_files = [n for n in zf.namelist() if n.startswith('word/media/')]

            for img_name in image_files:
                img_data = zf.read(img_name)
                ext = img_name.rsplit('.', 1)[-1].lower()
                mime = f"image/{ext}" if ext != 'jpg' else "image/jpeg"
                b64 = base64.b64encode(img_data).decode('utf-8')

                # Alle Bilder werden zunächst als Header-Bilder behandelt.
                # Das erste Bild ist in der Regel das Firmenlogo.
                result["header_images"].append({
                    "data": f"data:{mime};base64,{b64}",
                    "filename": img_name.split('/')[-1],
                })
    except Exception as e:
        logger.warning(f"Konnte Bilder nicht aus Vorlage extrahieren: {e}")

    return result
