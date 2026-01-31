#!/usr/bin/env python3
"""
Script zur Erstellung der Master-Templates für den Smart Document Generator.

Dieses Script erstellt die dedizierten Master-Templates pro Land und Kategorie
gemäß der v3.0 "Datei statt Logik" Strategie.

Usage:
    python scripts/create_master_templates.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Pfad zum Template-Verzeichnis
TEMPLATES_DIR = Path(__file__).parent.parent / "storage" / "master-templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def create_template_contract_de():
    """
    Erstellt das deutsche Vertrags-Template (DIN 5008).
    """
    doc = Document()

    # Seitenränder nach DIN 5008 Typ A
    section = doc.sections[0]
    section.left_margin = Mm(25)
    section.right_margin = Mm(20)
    section.top_margin = Mm(45)
    section.bottom_margin = Mm(25)

    # Header mit Platzhaltern
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = "{{ logo }}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_paragraph("{{ header_line1 }}")
    header.add_paragraph("{{ header_line2 }}")
    header.add_paragraph("{{ header_line3 }}")

    # Titel
    title = doc.add_heading("{{ document_title }}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertragsparteien
    doc.add_paragraph()
    doc.add_paragraph("zwischen")
    doc.add_paragraph()
    doc.add_paragraph("{{ company_name }}")
    doc.add_paragraph("{{ company_address }}")
    doc.add_paragraph("- nachfolgend 'Arbeitgeber' genannt -")
    doc.add_paragraph()
    doc.add_paragraph("und")
    doc.add_paragraph()
    doc.add_paragraph("{{ mitarbeiter_vorname }} {{ mitarbeiter_nachname }}")
    doc.add_paragraph("{{ mitarbeiter_adresse }}")
    doc.add_paragraph("- nachfolgend 'Arbeitnehmer' genannt -")
    doc.add_paragraph()

    # Platzhalter für assemblierten Content
    doc.add_paragraph("{{ assembled_content }}")

    # Unterschriften
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________          _______________________")
    doc.add_paragraph("Ort, Datum                               Ort, Datum")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________          _______________________")
    doc.add_paragraph("Arbeitgeber                               Arbeitnehmer")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Speichern
    output_path = TEMPLATES_DIR / "template_contract_de.docx"
    doc.save(output_path)
    print(f"✅ Erstellt: {output_path}")
    return output_path


def create_template_letter_de():
    """
    Erstellt das deutsche Brief-Template (DIN 5008 Geschäftsbrief).
    """
    doc = Document()

    # Seitenränder nach DIN 5008
    section = doc.sections[0]
    section.left_margin = Mm(25)
    section.right_margin = Mm(20)
    section.top_margin = Mm(27)
    section.bottom_margin = Mm(25)

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = "{{ logo }}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Absender
    doc.add_paragraph("{{ company_name }}")
    doc.add_paragraph("{{ company_address }}")
    doc.add_paragraph()

    # Empfänger
    doc.add_paragraph("{{ empfaenger_name }}")
    doc.add_paragraph("{{ empfaenger_adresse }}")
    doc.add_paragraph()

    # Datum
    date_para = doc.add_paragraph("{{ ort }}, {{ datum }}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    # Betreff
    betreff = doc.add_paragraph("{{ betreff }}")
    betreff.runs[0].bold = True if betreff.runs else None
    doc.add_paragraph()

    # Anrede
    doc.add_paragraph("{{ anrede }}")
    doc.add_paragraph()

    # Content
    doc.add_paragraph("{{ assembled_content }}")
    doc.add_paragraph()

    # Grußformel
    doc.add_paragraph("{{ grussformel }}")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{ unterschrift_name }}")
    doc.add_paragraph("{{ unterschrift_position }}")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = TEMPLATES_DIR / "template_letter_de.docx"
    doc.save(output_path)
    print(f"✅ Erstellt: {output_path}")
    return output_path


def create_template_contract_it():
    """
    Erstellt das italienische Vertrags-Template (UNI Norm).
    """
    doc = Document()

    # Italienische Seitenränder
    section = doc.sections[0]
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(25)

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = "{{ logo }}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph("{{ header_line1 }}")
    header.add_paragraph("{{ header_line2 }}")
    header.add_paragraph("{{ header_line3 }}")

    # Titel
    title = doc.add_heading("{{ document_title }}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Vertragsparteien (Italienisch)
    doc.add_paragraph()
    doc.add_paragraph("tra")
    doc.add_paragraph()
    doc.add_paragraph("{{ company_name }}")
    doc.add_paragraph("{{ company_address }}")
    doc.add_paragraph("- di seguito denominato 'Datore di lavoro' -")
    doc.add_paragraph()
    doc.add_paragraph("e")
    doc.add_paragraph()
    doc.add_paragraph("{{ dipendente_nome }} {{ dipendente_cognome }}")
    doc.add_paragraph("{{ dipendente_indirizzo }}")
    doc.add_paragraph("- di seguito denominato 'Lavoratore' -")
    doc.add_paragraph()

    # Content
    doc.add_paragraph("{{ assembled_content }}")

    # Unterschriften (Italienisch)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________          _______________________")
    doc.add_paragraph("Luogo, Data                              Luogo, Data")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________          _______________________")
    doc.add_paragraph("Datore di lavoro                         Lavoratore")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = TEMPLATES_DIR / "template_contract_it.docx"
    doc.save(output_path)
    print(f"✅ Erstellt: {output_path}")
    return output_path


def create_template_letter_it():
    """
    Erstellt das italienische Brief-Template.
    """
    doc = Document()

    # Italienische Seitenränder
    section = doc.sections[0]
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = "{{ logo }}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Absender
    doc.add_paragraph("{{ company_name }}")
    doc.add_paragraph("{{ company_address }}")
    doc.add_paragraph()

    # Empfänger
    doc.add_paragraph("Spett.le {{ destinatario_nome }}")
    doc.add_paragraph("{{ destinatario_indirizzo }}")
    doc.add_paragraph()

    # Datum (Italienisches Format)
    date_para = doc.add_paragraph("{{ luogo }}, {{ data }}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    # Oggetto (Betreff)
    oggetto = doc.add_paragraph("Oggetto: {{ oggetto }}")
    doc.add_paragraph()

    # Anrede
    doc.add_paragraph("{{ formula_apertura }}")
    doc.add_paragraph()

    # Content
    doc.add_paragraph("{{ assembled_content }}")
    doc.add_paragraph()

    # Grußformel
    doc.add_paragraph("{{ formula_chiusura }}")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{ firma_nome }}")
    doc.add_paragraph("{{ firma_posizione }}")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = TEMPLATES_DIR / "template_letter_it.docx"
    doc.save(output_path)
    print(f"✅ Erstellt: {output_path}")
    return output_path


def create_default_templates():
    """
    Erstellt Default-Templates als Fallback.
    """
    # Default DE
    doc_de = Document()
    section = doc_de.sections[0]
    section.left_margin = Mm(25)
    section.right_margin = Mm(20)
    section.top_margin = Mm(45)
    section.bottom_margin = Mm(25)

    doc_de.add_heading("{{ document_title }}", level=0)
    doc_de.add_paragraph()
    doc_de.add_paragraph("{{ assembled_content }}")

    footer_de = doc_de.sections[0].footer
    footer_para_de = footer_de.paragraphs[0] if footer_de.paragraphs else footer_de.add_paragraph()
    footer_para_de.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"

    output_de = TEMPLATES_DIR / "template_default_de.docx"
    doc_de.save(output_de)
    print(f"✅ Erstellt: {output_de}")

    # Default IT
    doc_it = Document()
    section_it = doc_it.sections[0]
    section_it.left_margin = Mm(20)
    section_it.right_margin = Mm(20)
    section_it.top_margin = Mm(30)
    section_it.bottom_margin = Mm(25)

    doc_it.add_heading("{{ document_title }}", level=0)
    doc_it.add_paragraph()
    doc_it.add_paragraph("{{ assembled_content }}")

    footer_it = doc_it.sections[0].footer
    footer_para_it = footer_it.paragraphs[0] if footer_it.paragraphs else footer_it.add_paragraph()
    footer_para_it.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"

    output_it = TEMPLATES_DIR / "template_default_it.docx"
    doc_it.save(output_it)
    print(f"✅ Erstellt: {output_it}")


def create_certificate_templates():
    """
    Erstellt Bescheinigungs-Templates.
    """
    for country, lang in [("de", "German"), ("it", "Italian")]:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Mm(25 if country == "de" else 20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(40)
        section.bottom_margin = Mm(25)

        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "{{ logo }}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_paragraph("{{ company_name }}")
        header.add_paragraph("{{ company_address }}")

        # Titel
        title = doc.add_heading("{{ document_title }}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph("{{ assembled_content }}")
        doc.add_paragraph()

        # Datum und Unterschrift
        if country == "de":
            doc.add_paragraph("{{ ort }}, den {{ datum }}")
        else:
            doc.add_paragraph("{{ luogo }}, {{ data }}")

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_______________________")
        doc.add_paragraph("{{ unterschrift_name }}" if country == "de" else "{{ firma_nome }}")

        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output_path = TEMPLATES_DIR / f"template_certificate_{country}.docx"
        doc.save(output_path)
        print(f"✅ Erstellt: {output_path}")


def create_amendment_templates():
    """
    Erstellt Vertragsänderungs-Templates.
    """
    for country, title_text in [("de", "Vertragsänderung / Nachtrag"), ("it", "Modifica contrattuale")]:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Mm(25 if country == "de" else 20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(45)
        section.bottom_margin = Mm(25)

        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "{{ logo }}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.add_paragraph("{{ header_line1 }}")
        header.add_paragraph("{{ header_line2 }}")

        # Titel
        title = doc.add_heading("{{ document_title }}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Bezug auf Ursprungsvertrag
        doc.add_paragraph()
        if country == "de":
            doc.add_paragraph("zum Arbeitsvertrag vom {{ ursprungsvertrag_datum }}")
        else:
            doc.add_paragraph("al contratto di lavoro del {{ data_contratto_originale }}")

        doc.add_paragraph()
        doc.add_paragraph("{{ assembled_content }}")
        doc.add_paragraph()

        # Unterschriften
        if country == "de":
            doc.add_paragraph("_______________________          _______________________")
            doc.add_paragraph("Arbeitgeber                               Arbeitnehmer")
        else:
            doc.add_paragraph("_______________________          _______________________")
            doc.add_paragraph("Datore di lavoro                         Lavoratore")

        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "{{ footer_line1 }} | {{ footer_line2 }} | {{ footer_line3 }}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output_path = TEMPLATES_DIR / f"template_amendment_{country}.docx"
        doc.save(output_path)
        print(f"✅ Erstellt: {output_path}")


def main():
    """
    Hauptfunktion: Erstellt alle Master-Templates.
    """
    print("=" * 60)
    print("Smart Document Generator - Master-Template Generator")
    print("=" * 60)
    print()
    print(f"Zielverzeichnis: {TEMPLATES_DIR}")
    print()

    # Alle Templates erstellen
    create_template_contract_de()
    create_template_letter_de()
    create_template_contract_it()
    create_template_letter_it()
    create_default_templates()
    create_certificate_templates()
    create_amendment_templates()

    print()
    print("=" * 60)
    print("✅ Alle Master-Templates wurden erfolgreich erstellt!")
    print("=" * 60)

    # Liste alle erstellten Templates
    print()
    print("Erstellte Templates:")
    for f in sorted(TEMPLATES_DIR.glob("*.docx")):
        print(f"  - {f.name}")


if __name__ == "__main__":
    main()
